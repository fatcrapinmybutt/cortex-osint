#!/usr/bin/env python3
"""THEMANBEARPIG v19.0 -- SINGULARITY CONVERGENCE.

Pinnacle build: 6 satellite engines (Athena, Automaton, Chronos, Cortex,
Oracle, Prometheus) fused with PROJECT KRAKEN + MBP Brain v5.0 + community
intelligence + pattern detection + filing pipeline into a single monolithic
litigation intelligence visualization.

Usage:
    python -I scripts/themanbearpig.py
    python -I scripts/themanbearpig.py --debug --hunt --export
"""

import argparse
import hashlib
import gzip
import http.server
import io
import json as _json
import mimetypes
import os
import random
import re
import shutil
import socket
import sqlite3
import subprocess
import sys
import threading
import time
from datetime import date, datetime
from pathlib import Path

# ---------------------------------------------------------------------------
# Satellite Engine Imports (lazy, graceful degradation)
# ---------------------------------------------------------------------------
HAS_ATHENA = False
HAS_AUTOMATON = False
HAS_CHRONOS = False
HAS_CORTEX = False
HAS_ORACLE = False
HAS_PROMETHEUS = False
HAS_BACKEND = False
HAS_BLEEDING_EDGE = False
HAS_INTELLIGENCE = False
HAS_OPERATIONS = False

try:
    from mbp_engines.athena import Athena
    HAS_ATHENA = True
except Exception:
    pass
try:
    from mbp_engines.automaton import AutomatonEngine
    HAS_AUTOMATON = True
except Exception:
    pass
try:
    from mbp_engines.chronos import ChronosEngine
    HAS_CHRONOS = True
except Exception:
    pass
try:
    from mbp_engines.cortex import CortexEngine
    HAS_CORTEX = True
except Exception:
    pass
try:
    from mbp_engines.oracle import OracleEngine
    HAS_ORACLE = True
except Exception:
    pass
try:
    from mbp_engines.prometheus import PrometheusEngine
    HAS_PROMETHEUS = True
except Exception:
    pass
try:
    from mbp_engines.backend_bridge import BackendBridge
    HAS_BACKEND = True
except Exception:
    pass
try:
    from mbp_engines.bleeding_edge_bridge import BleedingEdgeBridge
    HAS_BLEEDING_EDGE = True
except Exception:
    pass
try:
    from mbp_engines.intelligence_bridge import IntelligenceBridge
    HAS_INTELLIGENCE = True
except Exception:
    pass
try:
    from mbp_engines.operations_bridge import OperationsBridge
    HAS_OPERATIONS = True
except Exception:
    pass

# ---------------------------------------------------------------------------
# Paths -- PyInstaller-aware (sys._MEIPASS for bundled exe, normal for dev)
# ---------------------------------------------------------------------------
if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
    _BUNDLE = Path(sys._MEIPASS)
    REPO_ROOT = Path(sys.executable).resolve().parent
else:
    _BUNDLE = Path(__file__).resolve().parent.parent
    REPO_ROOT = _BUNDLE
BRAIN_DB = REPO_ROOT / "mbp_brain.db"
LIT_DB = REPO_ROOT / "litigation_context.db"
VIS_DIR = _BUNDLE / "08_MEDIA" / "MANBEARPIG_V15"
VIS_DIR_V9 = _BUNDLE / "08_MEDIA" / "MANBEARPIG_V9"
VIS_DIR_V5 = _BUNDLE / "08_MEDIA" / "MANBEARPIG_V5"
VIS_DIR_V7 = _BUNDLE / "08_MEDIA" / "MANBEARPIG_V7"
GRAPH_JSON = VIS_DIR / "graph_clusters.json"
if not GRAPH_JSON.exists():
    GRAPH_JSON = VIS_DIR_V9 / "graph_data.json"
if not GRAPH_JSON.exists():
    GRAPH_JSON = VIS_DIR_V5 / "graph_data.json"
EXPORT_SCRIPT = REPO_ROOT / "scripts" / "export_brain_d3.py"
EVOLVE_SCRIPT = REPO_ROOT / "scripts" / "brain_evolution.py"
KRAKEN_SCRIPT = REPO_ROOT / "07_CODE" / "PROJECT_KRAKEN" / "kraken.py"
FILING_SCRIPT = REPO_ROOT / "scripts" / "generate_filing.py"
COURT_FEED_SCRIPT = REPO_ROOT / "scripts" / "court_feed.py"
DOSSIER_DIR = REPO_ROOT / "04_ANALYSIS" / "ADVERSARY_TRACKS"

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
VERSION = "21.0.0"  # V7 SELFEVOLVE CONVERGENCE
SEPARATION_DATE = date(2025, 7, 29)
APP_BG = "#0a0a0f"
APP_WIDTH, APP_HEIGHT = 1920, 1080
APP_MIN_W, APP_MIN_H = 1280, 720

_PRAGMAS = (
    "PRAGMA busy_timeout = 60000",
    "PRAGMA journal_mode = WAL",
    "PRAGMA cache_size = -32000",
    "PRAGMA synchronous = NORMAL",
    "PRAGMA temp_store = MEMORY",
    "PRAGMA mmap_size = 268435456",  # 256 MB — brain DB is 310 MB on NVMe
)

# FTS5 sanitizer
_FTS_CLEAN = re.compile(r'[^\w\s*"]')

# Read-only guard
_WRITE_RE = re.compile(
    r"\b(INSERT|UPDATE|DELETE|DROP|CREATE|ALTER|REPLACE|ATTACH|DETACH|VACUUM|REINDEX)\b",
    re.IGNORECASE,
)


def _find_python():
    """Return path to the real Python interpreter, not the frozen exe.

    When running as a PyInstaller bundle, sys.executable points to the
    .exe itself.  Subprocess calls that need to run .py scripts must
    use the system Python instead.
    """
    if getattr(sys, "frozen", False):
        found = shutil.which("python")
        if found:
            return found
        # Fallback: common install location on this machine
        fallback = Path(r"C:\Users\andre\AppData\Local\Programs\Python\Python312\python.exe")
        if fallback.exists():
            return str(fallback)
        return "python"
    return sys.executable


_PYTHON = _find_python()

# Adversary patterns (from kraken.py)
ADVERSARIES = {
    "Emily Watson": r"(?i)\bemily\b.*\bwatson\b|\bwatson\b.*\bemily\b|\bemily\s+a\.?\s+watson\b",
    "Judge McNeill": r"(?i)\bmcneill\b|\bmcneil\b",
    "Pamela Rusco": r"(?i)\brusco\b|\bpamela\b.*\brusco\b",
    "Albert Watson": r"(?i)\balbert\b.*\bwatson\b|\bwatson\b.*\balbert\b",
    "Lori Watson": r"(?i)\blori\b.*\bwatson\b",
    "Ronald Berry": r"(?i)\bronald\b.*\bberry\b|\bron\b.*\bberry\b",
    "Cavan Berry": r"(?i)\bcavan\b.*\bberry\b",
    "Jennifer Barnes": r"(?i)\bbarnes\b.*\bjennifer\b|\bjennifer\b.*\bbarnes\b|\bP55406\b",
    "Kenneth Hoopes": r"(?i)\bhoopes\b|\bkenneth\b.*\bhoopes\b",
    "Maria Ladas-Hoopes": r"(?i)\bladas[\\s-]*hoopes\b|\bmaria\b.*\bladas\b",
    "FOC": r"(?i)\bfriend\s+of\s+(the\s+)?court\b|\bFOC\b",
    "Shady Oaks": r"(?i)\bshady\s*oaks\b",
}
ADVERSARY_RE = {k: re.compile(v) for k, v in ADVERSARIES.items()}

# Evidence categories
EVIDENCE_CATEGORIES = {
    "custody": re.compile(r"(?i)\bcustody\b|\bparenting\s*time\b|\bvisitation\b|\bbest\s+interest\b"),
    "PPO": re.compile(r"(?i)\bprotection\s*order\b|\bPPO\b|\brestraining\b|\bstalking\b"),
    "judicial": re.compile(r"(?i)\bjudicial\b|\bbias\b|\bex\s*parte\b|\brecusal\b|\bdisqualif"),
    "housing": re.compile(r"(?i)\beviction\b|\btenant\b|\blandlord\b|\bhousing\b|\bmobile\s*home\b"),
    "criminal": re.compile(r"(?i)\bcontempt\b|\bjail\b|\bincarcerat\b|\barrest\b|\bsentenc"),
    "financial": re.compile(r"(?i)\bchild\s*support\b|\bfiling\s*fee\b|\bdamages\b|\bgarnish"),
    "police": re.compile(r"(?i)\bpolice\b|\bNSPD\b|\bofficer\b|\bincident\s*report\b"),
    "medical": re.compile(r"(?i)\bhealthwest\b|\bmental\s*health\b|\bpsych\b|\bmedication\b|\bLOCUS\b"),
}

# Legal authority patterns
LEGAL_PATTERNS = {
    "MCR": re.compile(r"\bMCR\s+\d+\.\d+\w*"),
    "MCL": re.compile(r"\bMCL\s+\d+\.\d+\w*"),
    "MRE": re.compile(r"\bMRE\s+\d+\w*"),
    "USC": re.compile(r"\b\d+\s+U\.?S\.?C\.?\s*[S\s]*\d+"),
    "Case_Law": re.compile(
        r"\b\d+\s+Mich\.?\s+(App\.?\s+)?\d+|\b\d+\s+F\.\d[a-z]*\s+\d+"
    ),
}

# Focus mode boosters
FOCUS_BOOSTS = {
    "adversary": ["Emily Watson", "Albert Watson", "Lori Watson", "Ronald Berry", "Cavan Berry"],
    "judicial": ["Judge McNeill", "Kenneth Hoopes", "Maria Ladas-Hoopes", "Cavan Berry", "Pamela Rusco"],
    "housing": ["Shady Oaks"],
    "custody": ["Emily Watson", "FOC", "Pamela Rusco", "Albert Watson"],
    "ppo": ["Emily Watson", "Judge McNeill", "Ronald Berry"],
    "legal": [],
    "all": [],
}

EXTS = {".pdf", ".txt", ".csv", ".html", ".json", ".docx", ".md"}
MAX_CONTENT_BYTES = 500_000


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def _sanitize_fts(query: str) -> str:
    """Strip dangerous FTS5 metacharacters."""
    return _FTS_CLEAN.sub(" ", query).strip()


def _safe_float(val, default=0.0):
    """Coerce a value to float, returning *default* on any failure."""
    if val is None:
        return default
    try:
        return float(val)
    except (ValueError, TypeError):
        return default


def _sep_days() -> int:
    return (date.today() - SEPARATION_DATE).days


def _find_free_port() -> int:
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        s.bind(("127.0.0.1", 0))
        return s.getsockname()[1]


def _rows_to_dicts(rows):
    """Convert sqlite3.Row list to plain dicts for JSON serialization."""
    return [dict(r) for r in rows] if rows else []


def _connect_brain():
    """Open WAL-mode connection to mbp_brain.db."""
    if not BRAIN_DB.exists():
        return None
    conn = sqlite3.connect(str(BRAIN_DB))
    conn.row_factory = sqlite3.Row
    conn.text_factory = lambda b: b.decode("utf-8", errors="replace")
    for p in _PRAGMAS:
        conn.execute(p)
    return conn


def _connect_lit():
    """Open WAL-mode connection to litigation_context.db."""
    if not LIT_DB.exists():
        return None
    conn = sqlite3.connect(str(LIT_DB))
    conn.row_factory = sqlite3.Row
    conn.text_factory = lambda b: b.decode("utf-8", errors="replace")
    for p in _PRAGMAS:
        conn.execute(p)
    return conn


def _safe_query(conn, sql, params=(), limit=200):
    """Execute read-only SQL with automatic limit. Returns list of dicts."""
    if conn is None:
        return []
    try:
        rows = conn.execute(sql, params).fetchmany(limit)
        return _rows_to_dicts(rows)
    except Exception:
        return []


def _table_exists(conn, table_name):
    """Check if a table exists in the database."""
    if conn is None:
        return False
    try:
        row = conn.execute(
            "SELECT count(*) FROM sqlite_master WHERE type='table' AND name=?",
            (table_name,),
        ).fetchone()
        return row[0] > 0 if row else False
    except Exception:
        return False


def _fts_search(conn, fts_table, base_table, query, columns="*", limit=200):
    """FTS5 search with sanitization and LIKE fallback."""
    safe_q = _sanitize_fts(query)
    if not safe_q:
        return [], "empty"

    if _table_exists(conn, fts_table):
        try:
            sql = (
                f"SELECT {columns} FROM {fts_table} f "
                f"JOIN {base_table} n ON f.rowid = n.rowid "
                f"WHERE {fts_table} MATCH ? LIMIT ?"
            )
            rows = conn.execute(sql, (safe_q, limit)).fetchall()
            if rows:
                return _rows_to_dicts(rows), "fts5"
        except Exception:
            pass

    like_pat = f"%{safe_q}%"
    try:
        cols = [
            r[1]
            for r in conn.execute(f"PRAGMA table_info({base_table})").fetchall()
        ]
        text_cols = [c for c in cols if c in ("label", "description", "id", "quote_text", "event_description")]
        if not text_cols:
            text_cols = cols[:3]
        where = " OR ".join(f"{c} LIKE ?" for c in text_cols)
        sql = f"SELECT {columns} FROM {base_table} WHERE {where} LIMIT ?"
        params = tuple([like_pat] * len(text_cols)) + (limit,)
        rows = conn.execute(sql, params).fetchall()
        return _rows_to_dicts(rows), "like"
    except Exception:
        return [], "error"


# ---------------------------------------------------------------------------
# KRAKEN mini-engine (inline for thread-safe background hunting)
# ---------------------------------------------------------------------------
def _kraken_extract_text(filepath):
    """Extract text from a file (PDF/DOCX/TXT). Returns (text, method)."""
    ext = os.path.splitext(filepath)[1].lower()
    try:
        if ext == ".pdf":
            try:
                import pypdfium2 as pdfium
                pdf = pdfium.PdfDocument(filepath)
                pages = min(len(pdf), 30)
                text_parts = []
                for i in range(pages):
                    page = pdf[i]
                    tp = page.get_textpage()
                    text_parts.append(tp.get_text_range())
                    tp.close()
                    page.close()
                pdf.close()
                return "\n".join(text_parts)[:MAX_CONTENT_BYTES], f"PDF({pages}pp)"
            except ImportError:
                return "", "PDF_NO_LIB"
        elif ext == ".docx":
            try:
                from docx import Document
                doc = Document(filepath)
                text = "\n".join(p.text for p in doc.paragraphs)
                return text[:MAX_CONTENT_BYTES], f"DOCX({len(doc.paragraphs)}P)"
            except ImportError:
                return "", "DOCX_NO_LIB"
        elif ext in (".txt", ".csv", ".html", ".json", ".md"):
            for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
                try:
                    with open(filepath, "r", encoding=enc, errors="replace") as f:
                        return f.read(MAX_CONTENT_BYTES), f"TEXT({enc})"
                except Exception:
                    continue
            return "", "TEXT_ERR"
    except Exception as exc:
        return "", f"ERR({exc.__class__.__name__})"
    return "", "UNSUPPORTED"


def _kraken_analyze(content, filepath, focus="all"):
    """Analyze content for adversaries, legal authorities, categories, quotes."""
    if len(content) < 50:
        return {
            "adversaries": {},
            "legal": {},
            "categories": [],
            "key_quotes": [],
            "value_score": 0,
            "value_label": "EMPTY",
        }

    adversaries = {}
    for name, pat in ADVERSARY_RE.items():
        matches = pat.findall(content)
        if matches:
            adversaries[name] = len(matches)

    legal = {}
    for ltype, pat in LEGAL_PATTERNS.items():
        found = pat.findall(content)
        if found:
            legal[ltype] = list(set(found))

    categories = []
    for cat, pat in EVIDENCE_CATEGORIES.items():
        if pat.search(content):
            categories.append(cat)

    sentences = re.split(r"[.!?\n]+", content)
    key_quotes = []
    for sent in sentences:
        sent = sent.strip()
        if len(sent) < 30 or len(sent) > 500:
            continue
        for name, pat in ADVERSARY_RE.items():
            if name in adversaries and pat.search(sent):
                key_quotes.append(sent)
                break
        if len(key_quotes) >= 8:
            break

    adv_score = len(adversaries) * 3
    legal_score = len(legal) * 2
    cat_score = len(categories)
    quote_score = len(key_quotes)

    if focus in FOCUS_BOOSTS and FOCUS_BOOSTS[focus]:
        for name in FOCUS_BOOSTS[focus]:
            if name in adversaries:
                adv_score += adversaries[name]
    if focus == "legal":
        legal_score *= 2

    total = adv_score + legal_score + cat_score + quote_score
    if total >= 10:
        label = "HIGH"
    elif total >= 4:
        label = "MEDIUM"
    else:
        label = "LOW"

    return {
        "adversaries": adversaries,
        "legal": legal,
        "categories": categories,
        "key_quotes": key_quotes,
        "value_score": total,
        "value_label": label,
    }


def _kraken_discover_files(conn):
    """Discover candidate files from file_inventory DB table + local dirs."""
    files = []
    if conn and _table_exists(conn, "file_inventory"):
        try:
            rows = conn.execute(
                "SELECT file_path FROM file_inventory WHERE extension IN "
                "('.pdf','.txt','.csv','.html','.json','.docx','.md') "
                "ORDER BY RANDOM() LIMIT 5000"
            ).fetchall()
            files.extend(r[0] for r in rows if r[0])
        except Exception:
            pass

    scan_dirs = [
        REPO_ROOT / "01_EVIDENCE",
        REPO_ROOT / "02_AUTHORITY",
        REPO_ROOT / "04_ANALYSIS",
        REPO_ROOT / "05_FILINGS",
        REPO_ROOT / "09_REFERENCE",
    ]
    for d in scan_dirs:
        if d.exists():
            try:
                for root, _, fnames in os.walk(str(d)):
                    for fn in fnames:
                        if os.path.splitext(fn)[1].lower() in EXTS:
                            files.append(os.path.join(root, fn))
            except Exception:
                pass
    return list(set(files))


def _file_hash(fp):
    """Quick hash for dedup tracking."""
    try:
        st = os.stat(fp)
        raw = f"{fp}|{st.st_size}|{st.st_mtime}".encode()
        return hashlib.md5(raw).hexdigest()
    except Exception:
        return hashlib.md5(fp.encode()).hexdigest()


# ---------------------------------------------------------------------------
# UnifiedAPI -- JS bridge
# ---------------------------------------------------------------------------
class UnifiedAPI:
    """Exposes ALL methods to JavaScript via window.pywebview.api.*"""

    def __init__(self):
        self._brain_conn = None
        self._lit_conn = None
        self._kraken_thread = None
        self._kraken_status = {
            "running": False,
            "rounds_done": 0,
            "rounds_total": 0,
            "files_scanned": 0,
            "findings_high": 0,
            "findings_total": 0,
            "recent": [],
        }
        self._kraken_lock = threading.Lock()
        # Satellite engine instances (lazy)
        self._athena = None
        self._chronos = None
        self._cortex = None
        self._oracle = None
        self._prometheus = None
        self._automaton = None
        self._backend = None
        self._bleeding_edge = None
        self._intelligence = None
        self._operations = None

    # -- Connection helpers --

    def _brain(self):
        if self._brain_conn is None:
            self._brain_conn = _connect_brain()
        return self._brain_conn

    def _lit(self):
        if self._lit_conn is None:
            self._lit_conn = _connect_lit()
        return self._lit_conn

    # -- Satellite engine getters (lazy init, graceful degradation) --

    def _get_athena(self):
        if self._athena is None and HAS_ATHENA:
            try:
                self._athena = Athena(self._lit())
            except Exception:
                pass
        return self._athena

    def _get_chronos(self):
        if self._chronos is None and HAS_CHRONOS:
            try:
                self._chronos = ChronosEngine(self._lit())
            except Exception:
                pass
        return self._chronos

    def _get_cortex(self):
        if self._cortex is None and HAS_CORTEX:
            try:
                self._cortex = CortexEngine(self._lit())
            except Exception:
                pass
        return self._cortex

    def _get_oracle(self):
        if self._oracle is None and HAS_ORACLE:
            try:
                self._oracle = OracleEngine(self._lit())
            except Exception:
                pass
        return self._oracle

    def _get_prometheus(self):
        if self._prometheus is None and HAS_PROMETHEUS:
            try:
                self._prometheus = PrometheusEngine(self._lit())
            except Exception:
                pass
        return self._prometheus

    def _get_automaton(self):
        if self._automaton is None and HAS_AUTOMATON:
            try:
                self._automaton = AutomatonEngine(str(LIT_DB))
            except Exception:
                pass
        return self._automaton

    def _get_backend(self):
        if self._backend is None and HAS_BACKEND:
            try:
                self._backend = BackendBridge()
            except Exception:
                pass
        return self._backend

    def _get_bleeding_edge(self):
        if self._bleeding_edge is None and HAS_BLEEDING_EDGE:
            try:
                self._bleeding_edge = BleedingEdgeBridge(str(LIT_DB))
            except Exception:
                pass
        return self._bleeding_edge

    def _get_intelligence(self):
        if self._intelligence is None and HAS_INTELLIGENCE:
            try:
                self._intelligence = IntelligenceBridge(str(LIT_DB))
            except Exception:
                pass
        return self._intelligence

    def _get_operations(self):
        if self._operations is None and HAS_OPERATIONS:
            try:
                self._operations = OperationsBridge(str(LIT_DB))
            except Exception:
                pass
        return self._operations

    # ===================================================================
    # BRAIN API (from mbp_app.py BrainAPI)
    # ===================================================================

    def get_node_details(self, node_id):
        """Node + edges + chains + gaps for a given node ID."""
        conn = self._brain()
        if conn is None:
            return {"error": "Brain DB not available"}
        try:
            node = conn.execute("SELECT * FROM nodes WHERE id = ?", (node_id,)).fetchone()
            if not node:
                return {"error": f"Node {node_id} not found"}

            edges = conn.execute(
                "SELECT e.*, 'outgoing' AS direction FROM edges e WHERE e.source_id = ? "
                "UNION ALL "
                "SELECT e.*, 'incoming' AS direction FROM edges e WHERE e.target_id = ?",
                (node_id, node_id),
            ).fetchall()

            chains = conn.execute(
                "SELECT id, chain_path, strength_score, lane, filing_id "
                "FROM chains WHERE chain_path LIKE ? LIMIT 100",
                (f"%{node_id}%",),
            ).fetchall()

            gaps = []
            if _table_exists(conn, "gaps"):
                gaps = conn.execute(
                    "SELECT * FROM gaps WHERE node_id = ? AND resolved = 0",
                    (node_id,),
                ).fetchall()

            return {
                "node": dict(node),
                "edges": _rows_to_dicts(edges),
                "chains": _rows_to_dicts(chains),
                "gaps": _rows_to_dicts(gaps),
                "edge_count": len(edges),
                "chain_count": len(chains),
                "gap_count": len(gaps),
            }
        except Exception as exc:
            return {"error": str(exc)}

    def trace_chain(self, node_id):
        """All chains passing through a node."""
        conn = self._brain()
        if conn is None:
            return {"error": "Brain DB not available", "chains": [], "count": 0}
        try:
            rows = conn.execute(
                "SELECT * FROM chains WHERE chain_path LIKE ? "
                "ORDER BY strength_score DESC LIMIT 200",
                (f"%{node_id}%",),
            ).fetchall()
            chains = []
            for r in rows:
                d = dict(r)
                try:
                    d["path_nodes"] = _json.loads(d.get("chain_path", "[]"))
                except Exception:
                    d["path_nodes"] = []
                chains.append(d)
            return {"chains": chains, "count": len(chains)}
        except Exception as exc:
            return {"error": str(exc), "chains": [], "count": 0}

    def search_nodes(self, query):
        """FTS5 search across brain nodes with LIKE fallback."""
        if not query or not query.strip():
            return {"results": [], "count": 0, "method": "empty"}
        conn = self._brain()
        if conn is None:
            return {"results": [], "count": 0, "method": "no_db"}
        results, method = _fts_search(conn, "nodes_fts", "nodes", query, "n.*")
        if method == "fts5":
            return {"results": results, "count": len(results), "method": "fts5"}

        safe_q = _sanitize_fts(query)
        if not safe_q:
            return {"results": [], "count": 0, "method": "sanitized_empty"}
        like_pat = f"%{safe_q}%"
        try:
            rows = conn.execute(
                "SELECT * FROM nodes WHERE label LIKE ? OR description LIKE ? OR id LIKE ? "
                "ORDER BY CASE WHEN label LIKE ? THEN 0 ELSE 1 END LIMIT 200",
                (like_pat, like_pat, like_pat, like_pat),
            ).fetchall()
            return {"results": _rows_to_dicts(rows), "count": len(rows), "method": "like"}
        except Exception as exc:
            return {"results": [], "count": 0, "method": "error", "error": str(exc)}

    def get_stats(self):
        """Brain stats + separation days."""
        conn = self._brain()
        result = {"separation_days": _sep_days(), "brain_available": conn is not None}
        if conn is None:
            return result
        try:
            row = conn.execute(
                "SELECT "
                "(SELECT COUNT(*) FROM nodes) AS node_count, "
                "(SELECT COUNT(*) FROM edges) AS edge_count, "
                "(SELECT COUNT(*) FROM chains) AS chain_count, "
                "(SELECT COUNT(*) FROM gaps) AS gap_total, "
                "(SELECT COUNT(*) FROM gaps WHERE resolved = 0) AS gap_open, "
                "(SELECT COUNT(*) FROM gaps WHERE priority = 'HIGH' AND resolved = 0) AS gap_high, "
                "(SELECT MAX(version) FROM versions) AS brain_version"
            ).fetchone()
            result.update(dict(row))

            layers = conn.execute(
                "SELECT layer, COUNT(*) AS cnt FROM nodes GROUP BY layer ORDER BY cnt DESC"
            ).fetchall()
            result["layers"] = _rows_to_dicts(layers)

            edge_types = conn.execute(
                "SELECT edge_type, COUNT(*) AS cnt FROM edges GROUP BY edge_type ORDER BY cnt DESC"
            ).fetchall()
            result["edge_types"] = _rows_to_dicts(edge_types)

            lanes = conn.execute(
                "SELECT lane, COUNT(*) AS cnt FROM nodes WHERE lane != '' "
                "GROUP BY lane ORDER BY cnt DESC"
            ).fetchall()
            result["lanes"] = _rows_to_dicts(lanes)

            top = conn.execute(
                "SELECT id, strength_score, lane, filing_id FROM chains "
                "ORDER BY strength_score DESC LIMIT 5"
            ).fetchall()
            result["top_chains"] = _rows_to_dicts(top)
        except Exception as exc:
            result["error"] = str(exc)
        return result

    def get_gaps(self, priority=""):
        """Gap listing, optionally filtered by priority."""
        conn = self._brain()
        if conn is None:
            return {"gaps": [], "count": 0, "by_priority": {}}
        try:
            if priority and priority.upper() in ("HIGH", "MEDIUM", "LOW"):
                rows = conn.execute(
                    "SELECT * FROM gaps WHERE resolved = 0 AND priority = ? "
                    "ORDER BY created_at DESC LIMIT 500",
                    (priority.upper(),),
                ).fetchall()
            else:
                rows = conn.execute(
                    "SELECT * FROM gaps WHERE resolved = 0 "
                    "ORDER BY CASE priority WHEN 'HIGH' THEN 0 WHEN 'MEDIUM' THEN 1 ELSE 2 END, "
                    "created_at DESC LIMIT 500"
                ).fetchall()

            gaps = _rows_to_dicts(rows)
            by_p = {"HIGH": 0, "MEDIUM": 0, "LOW": 0, "UNKNOWN": 0}
            for g in gaps:
                p = g.get("priority", "UNKNOWN")
                by_p[p] = by_p.get(p, 0) + 1
            return {"gaps": gaps, "count": len(gaps), "by_priority": by_p}
        except Exception as exc:
            return {"gaps": [], "count": 0, "error": str(exc), "by_priority": {}}

    def run_evolution(self):
        """Trigger brain_evolution.py subprocess."""
        if not EVOLVE_SCRIPT.exists():
            return {"error": f"Evolution script not found: {EVOLVE_SCRIPT}"}
        try:
            result = subprocess.run(
                [_PYTHON, "-I", str(EVOLVE_SCRIPT), "--stats"],
                capture_output=True, text=True, timeout=120,
                cwd=str(REPO_ROOT),
            )
            return {
                "stdout": result.stdout[-4000:],
                "stderr": result.stderr[-2000:],
                "returncode": result.returncode,
            }
        except subprocess.TimeoutExpired:
            return {"error": "Evolution script timed out after 120 seconds"}
        except Exception as exc:
            return {"error": str(exc)}

    def refresh_data(self):
        """Re-export graph_data.json from brain DB."""
        if not EXPORT_SCRIPT.exists():
            return {"error": f"Export script not found: {EXPORT_SCRIPT}"}
        try:
            result = subprocess.run(
                [_PYTHON, "-I", str(EXPORT_SCRIPT)],
                capture_output=True, text=True, timeout=120,
                cwd=str(REPO_ROOT),
            )
            return {
                "success": result.returncode == 0,
                "stdout": result.stdout[-4000:],
                "stderr": result.stderr[-2000:],
                "returncode": result.returncode,
                "message": (
                    "Reload the page to see updated data"
                    if result.returncode == 0
                    else "Export failed -- check stderr"
                ),
            }
        except subprocess.TimeoutExpired:
            return {"error": "Export script timed out after 120 seconds"}
        except Exception as exc:
            return {"error": str(exc)}

    # ===================================================================
    # LITIGATION API (from adversary_blueprint.py LitigationAPI)
    # ===================================================================

    def get_live_data(self):
        """Separation counter + DB stats."""
        conn = self._lit()
        stats = {"ev": 0, "jv": 0, "im": 0, "tl": 0, "ct": 0}
        if conn:
            try:
                row = conn.execute(
                    "SELECT "
                    "(SELECT COUNT(*) FROM evidence_quotes) AS ev, "
                    "(SELECT COUNT(*) FROM judicial_violations) AS jv, "
                    "(SELECT COUNT(*) FROM impeachment_matrix) AS im, "
                    "(SELECT COUNT(*) FROM timeline_events) AS tl, "
                    "(SELECT COUNT(*) FROM contradiction_map) AS ct"
                ).fetchone()
                stats = dict(row)
            except Exception:
                pass
        return {
            "sep_days": _sep_days(),
            "sep_date": str(SEPARATION_DATE),
            "today": str(date.today()),
            "db_available": conn is not None,
            "stats": stats,
        }

    def query_evidence(self, query, limit=20):
        """Evidence quote search from litigation_context.db."""
        conn = self._lit()
        if conn is None:
            return []
        safe_q = _sanitize_fts(query)
        if not safe_q:
            return []

        if _table_exists(conn, "evidence_fts"):
            try:
                rows = conn.execute(
                    "SELECT eq.id, eq.quote_text, eq.source_file, eq.page_number, "
                    "eq.category, eq.lane, eq.relevance_score "
                    "FROM evidence_fts f JOIN evidence_quotes eq ON f.rowid = eq.rowid "
                    "WHERE evidence_fts MATCH ? LIMIT ?",
                    (safe_q, limit),
                ).fetchall()
                if rows:
                    return _rows_to_dicts(rows)
            except Exception:
                pass

        like_pat = f"%{safe_q}%"
        return _safe_query(
            conn,
            "SELECT id, quote_text, source_file, page_number, category, lane, relevance_score "
            "FROM evidence_quotes WHERE quote_text LIKE ? OR source_file LIKE ? "
            "ORDER BY relevance_score DESC LIMIT ?",
            (like_pat, like_pat, limit),
            limit,
        )

    def query_adversary(self, name):
        """Adversary intelligence dossier."""
        conn = self._lit()
        if conn is None:
            return {"name": name, "error": "DB not available"}
        like = f"%{name}%"
        result = {"name": name, "evidence_count": 0, "violations": [], "impeachment": [], "contradictions": []}
        try:
            ev = conn.execute(
                "SELECT COUNT(*) FROM evidence_quotes WHERE quote_text LIKE ?",
                (like,),
            ).fetchone()
            result["evidence_count"] = ev[0] if ev else 0

            if _table_exists(conn, "judicial_violations"):
                result["violations"] = _safe_query(
                    conn,
                    "SELECT violation_type, COUNT(*) AS cnt FROM judicial_violations "
                    "WHERE description LIKE ? GROUP BY violation_type ORDER BY cnt DESC LIMIT 20",
                    (like,),
                )

            if _table_exists(conn, "impeachment_matrix"):
                result["impeachment"] = _safe_query(
                    conn,
                    "SELECT category, COUNT(*) AS cnt, AVG(impeachment_value) AS avg_sev "
                    "FROM impeachment_matrix WHERE target LIKE ? "
                    "GROUP BY category ORDER BY cnt DESC LIMIT 20",
                    (like,),
                )

            if _table_exists(conn, "contradiction_map"):
                result["contradictions"] = _safe_query(
                    conn,
                    "SELECT actor, statement_1, statement_2, severity, source_1, source_2 "
                    "FROM contradiction_map WHERE actor LIKE ? "
                    "ORDER BY severity DESC LIMIT 20",
                    (like,),
                )
        except Exception as exc:
            result["error"] = str(exc)
        return result

    def query_timeline(self, start="", end="", limit=50):
        """Timeline events with date range filter."""
        conn = self._lit()
        if conn is None:
            return []
        conditions = []
        params = []
        if start:
            conditions.append("event_date >= ?")
            params.append(start)
        if end:
            conditions.append("event_date <= ?")
            params.append(end)
        where = " AND ".join(conditions) if conditions else "1=1"
        params.append(limit)
        return _safe_query(
            conn,
            f"SELECT event_date, event_description, actors, source_document "
            f"FROM timeline_events WHERE {where} ORDER BY event_date DESC LIMIT ?",
            tuple(params),
            limit,
        )

    def query_judicial(self, jtype=""):
        """Judicial violations by type."""
        conn = self._lit()
        if conn is None:
            return {"violations": [], "cartel": []}
        result = {"violations": [], "cartel": []}
        try:
            if jtype:
                result["violations"] = _safe_query(
                    conn,
                    "SELECT violation_type, description, COUNT(*) AS cnt "
                    "FROM judicial_violations WHERE violation_type LIKE ? "
                    "GROUP BY violation_type ORDER BY cnt DESC LIMIT 30",
                    (f"%{jtype}%",),
                )
            else:
                result["violations"] = _safe_query(
                    conn,
                    "SELECT violation_type, COUNT(*) AS cnt "
                    "FROM judicial_violations GROUP BY violation_type ORDER BY cnt DESC LIMIT 30",
                )

            if _table_exists(conn, "berry_mcneill_intelligence"):
                result["cartel"] = _safe_query(
                    conn,
                    "SELECT person, connection_type, description, evidence_source "
                    "FROM berry_mcneill_intelligence ORDER BY rowid DESC LIMIT 50",
                )
        except Exception as exc:
            result["error"] = str(exc)
        return result

    def query_impeachment(self, target):
        """Impeachment matrix for a target."""
        conn = self._lit()
        if conn is None:
            return []
        return _safe_query(
            conn,
            "SELECT target, category, evidence_summary, impeachment_value, "
            "cross_exam_question, source_file FROM impeachment_matrix "
            "WHERE target LIKE ? ORDER BY impeachment_value DESC LIMIT 50",
            (f"%{target}%",),
        )

    def search_fts5(self, query):
        """Full-text search across multiple tables."""
        conn = self._lit()
        if conn is None:
            return {"results": [], "count": 0}
        safe_q = _sanitize_fts(query)
        if not safe_q:
            return {"results": [], "count": 0}
        all_results = []
        tables = [
            ("evidence_fts", "evidence_quotes", "quote_text", "evidence"),
            ("timeline_fts", "timeline_events", "event_description", "timeline"),
        ]
        for fts_t, base_t, text_col, source_tag in tables:
            if _table_exists(conn, fts_t):
                try:
                    rows = conn.execute(
                        f"SELECT b.rowid, b.{text_col} AS text FROM {fts_t} f "
                        f"JOIN {base_t} b ON f.rowid = b.rowid "
                        f"WHERE {fts_t} MATCH ? LIMIT 50",
                        (safe_q,),
                    ).fetchall()
                    for r in rows:
                        all_results.append({"source": source_tag, "text": r["text"], "rowid": r["rowid"]})
                except Exception:
                    pass

        if not all_results:
            like_pat = f"%{safe_q}%"
            for base_t, text_col, tag in [
                ("evidence_quotes", "quote_text", "evidence"),
                ("timeline_events", "event_description", "timeline"),
            ]:
                if _table_exists(conn, base_t):
                    try:
                        rows = conn.execute(
                            f"SELECT rowid, {text_col} AS text FROM {base_t} "
                            f"WHERE {text_col} LIKE ? LIMIT 30",
                            (like_pat,),
                        ).fetchall()
                        for r in rows:
                            all_results.append({"source": tag, "text": r["text"], "rowid": r["rowid"]})
                    except Exception:
                        pass

        return {"results": all_results, "count": len(all_results)}

    # ===================================================================
    # KRAKEN HUNTING (background thread)
    # ===================================================================

    def start_kraken(self, rounds=3, count=10, focus="all"):
        """Launch KRAKEN hunting in background thread."""
        with self._kraken_lock:
            if self._kraken_status["running"]:
                return {"status": "already_running", "progress": self._kraken_status}

        self._kraken_status = {
            "running": True,
            "rounds_done": 0,
            "rounds_total": int(rounds),
            "files_scanned": 0,
            "findings_high": 0,
            "findings_total": 0,
            "recent": [],
        }

        def _hunt():
            lit_conn = _connect_lit()
            all_files = _kraken_discover_files(lit_conn)
            processed_hashes = set()

            if lit_conn and _table_exists(lit_conn, "kraken_processed"):
                try:
                    rows = lit_conn.execute("SELECT file_hash FROM kraken_processed").fetchall()
                    processed_hashes = {r[0] for r in rows}
                except Exception:
                    pass

            for rnd in range(int(rounds)):
                candidates = []
                for fp in all_files:
                    fh = _file_hash(fp)
                    if fh not in processed_hashes and os.path.isfile(fp):
                        try:
                            sz = os.path.getsize(fp)
                            if 100 <= sz <= 50_000_000:
                                candidates.append((fp, fh, sz))
                        except Exception:
                            pass

                sample = random.sample(candidates, min(int(count), len(candidates))) if candidates else []
                for fp, fh, sz in sample:
                    text, method = _kraken_extract_text(fp)
                    analysis = _kraken_analyze(text, fp, focus)
                    processed_hashes.add(fh)

                    with self._kraken_lock:
                        self._kraken_status["files_scanned"] += 1
                        self._kraken_status["findings_total"] += 1
                        if analysis["value_label"] == "HIGH":
                            self._kraken_status["findings_high"] += 1

                        if analysis["value_score"] >= 4:
                            entry = {
                                "file": os.path.basename(fp),
                                "score": analysis["value_score"],
                                "label": analysis["value_label"],
                                "adversaries": list(analysis["adversaries"].keys()),
                                "categories": analysis["categories"],
                                "quotes": analysis["key_quotes"][:2],
                            }
                            recent = self._kraken_status["recent"]
                            recent.insert(0, entry)
                            self._kraken_status["recent"] = recent[:50]

                    if analysis["value_label"] == "HIGH" and lit_conn:
                        for quote in analysis["key_quotes"][:5]:
                            try:
                                cats = ",".join(analysis["categories"][:3])
                                lane = "A"
                                if "judicial" in analysis["categories"]:
                                    lane = "E"
                                elif "housing" in analysis["categories"]:
                                    lane = "B"
                                elif "PPO" in analysis["categories"]:
                                    lane = "D"
                                elif "criminal" in analysis["categories"]:
                                    lane = "CRIMINAL"
                                lit_conn.execute(
                                    "INSERT OR IGNORE INTO evidence_quotes "
                                    "(source_file, quote_text, category, lane, relevance_score, created_at) "
                                    "VALUES (?, ?, ?, ?, ?, datetime('now'))",
                                    (os.path.basename(fp), quote, cats, lane, analysis["value_score"]),
                                )
                            except Exception:
                                pass
                        try:
                            lit_conn.commit()
                        except Exception:
                            pass

                    if lit_conn:
                        try:
                            if not _table_exists(lit_conn, "kraken_processed"):
                                lit_conn.execute(
                                    "CREATE TABLE IF NOT EXISTS kraken_processed ("
                                    "file_hash TEXT PRIMARY KEY, file_path TEXT, "
                                    "processed_at TEXT DEFAULT (datetime('now')), "
                                    "value_score INTEGER DEFAULT 0, value_label TEXT DEFAULT 'LOW')"
                                )
                            lit_conn.execute(
                                "INSERT OR IGNORE INTO kraken_processed (file_hash, file_path, value_score, value_label) "
                                "VALUES (?, ?, ?, ?)",
                                (fh, fp, analysis["value_score"], analysis["value_label"]),
                            )
                            lit_conn.commit()
                        except Exception:
                            pass

                with self._kraken_lock:
                    self._kraken_status["rounds_done"] = rnd + 1

            if lit_conn:
                try:
                    lit_conn.close()
                except Exception:
                    pass
            with self._kraken_lock:
                self._kraken_status["running"] = False

        t = threading.Thread(target=_hunt, daemon=True, name="kraken-hunter")
        t.start()
        self._kraken_thread = t
        return {"status": "started", "rounds": int(rounds), "count": int(count), "focus": focus}

    def get_kraken_status(self):
        """Current hunting progress."""
        with self._kraken_lock:
            return dict(self._kraken_status)

    # ===================================================================
    # FILING GENERATION (from generate_filing.py)
    # ===================================================================

    def list_filings(self):
        """Available filing IDs with chain statistics."""
        conn = self._brain()
        if conn is None:
            return []
        try:
            rows = conn.execute(
                "SELECT filing_id, COUNT(*) AS chain_count, "
                "AVG(strength_score) AS avg_strength, MAX(strength_score) AS max_strength "
                "FROM chains WHERE filing_id IS NOT NULL AND filing_id != '' "
                "GROUP BY filing_id ORDER BY avg_strength DESC"
            ).fetchall()
            results = []
            for r in rows:
                d = dict(r)
                node = conn.execute(
                    "SELECT label, lane FROM nodes WHERE id = ? LIMIT 1",
                    (d["filing_id"],),
                ).fetchone()
                if node:
                    d["label"] = node["label"]
                    d["lane"] = node["lane"]
                else:
                    d["label"] = d["filing_id"]
                    d["lane"] = ""
                results.append(d)
            return results
        except Exception as exc:
            return [{"error": str(exc)}]

    def generate_brief(self, filing_id):
        """Auto-generate a brief for a filing via generate_filing.py."""
        if not FILING_SCRIPT.exists():
            return {"error": "Filing script not found"}
        try:
            result = subprocess.run(
                [_PYTHON, "-I", str(FILING_SCRIPT), "--filing", filing_id, "--brief"],
                capture_output=True, text=True, timeout=60,
                cwd=str(REPO_ROOT),
            )
            return {
                "text": result.stdout[-16000:] if result.returncode == 0 else "",
                "stderr": result.stderr[-2000:],
                "returncode": result.returncode,
                "success": result.returncode == 0,
            }
        except subprocess.TimeoutExpired:
            return {"error": "Brief generation timed out (60s)"}
        except Exception as exc:
            return {"error": str(exc)}

    def generate_impeachment(self, actor_id):
        """Cross-examination outline for an actor."""
        if not FILING_SCRIPT.exists():
            return {"error": "Filing script not found"}
        try:
            result = subprocess.run(
                [_PYTHON, "-I", str(FILING_SCRIPT), "--actor", actor_id, "--impeach"],
                capture_output=True, text=True, timeout=60,
                cwd=str(REPO_ROOT),
            )
            return {
                "text": result.stdout[-16000:] if result.returncode == 0 else "",
                "stderr": result.stderr[-2000:],
                "returncode": result.returncode,
                "success": result.returncode == 0,
            }
        except subprocess.TimeoutExpired:
            return {"error": "Impeachment generation timed out (60s)"}
        except Exception as exc:
            return {"error": str(exc)}

    def get_strongest_filing(self):
        """Strongest chain -> filing info."""
        conn = self._brain()
        if conn is None:
            return {"error": "Brain DB not available"}
        try:
            row = conn.execute(
                "SELECT filing_id, MAX(strength_score) AS max_str, lane "
                "FROM chains WHERE filing_id IS NOT NULL AND filing_id != '' "
                "GROUP BY filing_id ORDER BY max_str DESC LIMIT 1"
            ).fetchone()
            if row:
                d = dict(row)
                node = conn.execute(
                    "SELECT label, description FROM nodes WHERE id = ?",
                    (d["filing_id"],),
                ).fetchone()
                if node:
                    d["label"] = node["label"]
                    d["description"] = node["description"]
                return d
            return {"error": "No filings found"}
        except Exception as exc:
            return {"error": str(exc)}

    # ===================================================================
    # COURT FEED
    # ===================================================================

    def fetch_court_feed(self):
        """Fetch latest Michigan court updates."""
        if not COURT_FEED_SCRIPT.exists():
            return {"error": "court_feed.py not found", "items": []}
        try:
            result = subprocess.run(
                [_PYTHON, "-I", str(COURT_FEED_SCRIPT), "--json"],
                capture_output=True, text=True, timeout=30,
                cwd=str(REPO_ROOT),
            )
            if result.returncode == 0 and result.stdout.strip():
                try:
                    return _json.loads(result.stdout)
                except Exception:
                    return {"items": [], "raw": result.stdout[:2000]}
            return {"items": [], "stderr": result.stderr[:1000]}
        except subprocess.TimeoutExpired:
            return {"error": "Court feed timed out (30s)", "items": []}
        except Exception as exc:
            return {"error": str(exc), "items": []}

    # ===================================================================
    # ANALYTICS PASSTHROUGH
    # ===================================================================

    def run_analytics(self, command):
        """Run an mbp_intel.py analytics command."""
        intel_script = REPO_ROOT / "07_CODE" / "PROJECT_KRAKEN" / "mbp" / "mbp_intel.py"
        if not intel_script.exists():
            return {"error": "mbp_intel.py not found"}
        if _WRITE_RE.search(command):
            return {"error": "Write operations not allowed"}
        try:
            result = subprocess.run(
                [_PYTHON, "-I", str(intel_script), command],
                capture_output=True, text=True, timeout=30,
                cwd=str(REPO_ROOT),
            )
            return {
                "output": result.stdout[-8000:],
                "stderr": result.stderr[-2000:],
                "returncode": result.returncode,
            }
        except subprocess.TimeoutExpired:
            return {"error": f"Analytics command '{command}' timed out (30s)"}
        except Exception as exc:
            return {"error": str(exc)}

    # ===================================================================
    # CHRONOLOGY / FILING / ADVERSARY / DASHBOARD / EXPORT / SEARCH
    # ===================================================================

    def get_chronology(self, lane=None, limit=200):
        """Timeline events from litigation_context.db, ordered by date.

        Args:
            lane: Optional lane filter (e.g. "A", "D").
            limit: Max rows to return (default 200).
        Returns:
            dict with ``events`` list and ``count``.
        """
        conn = self._lit()
        if conn is None:
            return {"events": [], "count": 0, "error": "Litigation DB not available"}
        try:
            if lane:
                sql = (
                    "SELECT id, event_date, event_description, actors, lane, "
                    "category, source_table, source_id, severity, filing_relevance "
                    "FROM timeline_events WHERE lane = ? "
                    "ORDER BY event_date ASC LIMIT ?"
                )
                rows = conn.execute(sql, (lane, limit)).fetchall()
            else:
                sql = (
                    "SELECT id, event_date, event_description, actors, lane, "
                    "category, source_table, source_id, severity, filing_relevance "
                    "FROM timeline_events ORDER BY event_date ASC LIMIT ?"
                )
                rows = conn.execute(sql, (limit,)).fetchall()

            events = []
            for r in rows:
                d = dict(r)
                d["severity"] = _safe_float(d.get("severity"))
                events.append(d)

            return {"events": events, "count": len(events)}
        except Exception as exc:
            return {"events": [], "count": 0, "error": str(exc)}

    def get_filing_packet(self, filing_id):
        """Comprehensive packet overview for a filing from mbp_brain.db.

        Args:
            filing_id: e.g. ``"FILING_MCR_2003_DISQUALIFICATION"``.
        Returns:
            dict with filing details, chains, evidence/violation/authority
            counts, exhibit candidates, and blocking gaps.
        """
        conn = self._brain()
        if conn is None:
            return {"error": "Brain DB not available"}
        if not filing_id:
            return {"error": "filing_id is required"}
        try:
            # Filing node
            node = conn.execute(
                "SELECT * FROM nodes WHERE id = ?", (filing_id,)
            ).fetchone()
            if not node:
                return {"error": f"Filing node '{filing_id}' not found"}
            filing_info = dict(node)
            filing_info["readiness"] = _safe_float(filing_info.get("readiness"))
            filing_info["severity"] = _safe_float(filing_info.get("severity"))
            filing_info["confidence"] = _safe_float(filing_info.get("confidence"))

            # Chains reaching this filing, sorted by strength
            chain_rows = conn.execute(
                "SELECT id, chain_path, chain_type, total_weight, length, "
                "lane, filing_id, evidence_ids, strength_score "
                "FROM chains WHERE filing_id = ? "
                "ORDER BY strength_score DESC LIMIT 200",
                (filing_id,),
            ).fetchall()
            chains = []
            for r in chain_rows:
                d = dict(r)
                d["strength_score"] = _safe_float(d.get("strength_score"))
                d["total_weight"] = _safe_float(d.get("total_weight"))
                chains.append(d)

            # Count supporting nodes by type across all chains
            evidence_count = 0
            violation_count = 0
            authority_count = 0
            exhibit_candidates = []
            seen_evidence = set()

            for ch in chains:
                raw_ids = ch.get("evidence_ids") or "[]"
                try:
                    eid_list = _json.loads(raw_ids) if raw_ids.startswith("[") else []
                except Exception:
                    eid_list = []
                for eid in eid_list:
                    if eid not in seen_evidence:
                        seen_evidence.add(eid)

            if seen_evidence:
                placeholders = ",".join("?" for _ in seen_evidence)
                type_rows = conn.execute(
                    f"SELECT node_type, COUNT(*) AS cnt FROM nodes "
                    f"WHERE id IN ({placeholders}) GROUP BY node_type",
                    list(seen_evidence),
                ).fetchall()
                for tr in type_rows:
                    nt = (tr["node_type"] or "").lower()
                    cnt = tr["cnt"]
                    if "evidence" in nt or "quote" in nt:
                        evidence_count += cnt
                    elif "violation" in nt or "judicial" in nt:
                        violation_count += cnt
                    elif "authority" in nt or "rule" in nt or "statute" in nt:
                        authority_count += cnt

                # Exhibit candidates = evidence-type nodes in chains
                ex_rows = conn.execute(
                    f"SELECT id, label, node_type, lane, confidence FROM nodes "
                    f"WHERE id IN ({placeholders}) "
                    f"AND (node_type LIKE '%evidence%' OR node_type LIKE '%quote%' "
                    f"     OR node_type LIKE '%exhibit%' OR node_type LIKE '%police%')",
                    list(seen_evidence),
                ).fetchall()
                for er in ex_rows:
                    d = dict(er)
                    d["confidence"] = _safe_float(d.get("confidence"))
                    exhibit_candidates.append(d)

            # Gaps blocking this filing
            gaps = []
            if _table_exists(conn, "gaps"):
                gap_rows = conn.execute(
                    "SELECT id, gap_type, description, priority, acquisition_task "
                    "FROM gaps WHERE node_id = ? AND resolved = 0 "
                    "ORDER BY priority DESC",
                    (filing_id,),
                ).fetchall()
                gaps = _rows_to_dicts(gap_rows)

            return {
                "filing": filing_info,
                "chains": chains,
                "chain_count": len(chains),
                "evidence_count": evidence_count,
                "violation_count": violation_count,
                "authority_count": authority_count,
                "exhibit_candidates": exhibit_candidates,
                "gaps": gaps,
                "gap_count": len(gaps),
            }
        except Exception as exc:
            return {"error": str(exc)}

    def get_adversary_overview(self):
        """All adversary analytics in a single batch query.

        Returns a dict keyed by adversary name, each containing
        evidence_count, violation_count, impeachment_count, and
        top_contradiction from litigation_context.db.
        """
        conn = self._lit()
        if conn is None:
            return {"adversaries": {}, "error": "Litigation DB not available"}
        try:
            adversary_names = list(ADVERSARIES.keys())
            result = {}
            for name in adversary_names:
                result[name] = {
                    "evidence_count": 0,
                    "violation_count": 0,
                    "impeachment_count": 0,
                    "top_contradiction": None,
                }

            # Evidence counts -- batch via CASE WHEN
            case_clauses = []
            params = []
            for name in adversary_names:
                case_clauses.append(
                    "SUM(CASE WHEN quote_text LIKE ? OR source_file LIKE ? THEN 1 ELSE 0 END)"
                )
                like = f"%{name.split()[0]}%"
                params.extend([like, like])
            sql_ev = f"SELECT {', '.join(case_clauses)} FROM evidence_quotes WHERE is_duplicate = 0"
            row = conn.execute(sql_ev, params).fetchone()
            if row:
                for i, name in enumerate(adversary_names):
                    result[name]["evidence_count"] = row[i] or 0

            # Judicial violation counts per adversary
            jv_clauses = []
            jv_params = []
            for name in adversary_names:
                jv_clauses.append(
                    "SUM(CASE WHEN description LIKE ? OR source_quote LIKE ? THEN 1 ELSE 0 END)"
                )
                like = f"%{name.split()[0]}%"
                jv_params.extend([like, like])
            sql_jv = f"SELECT {', '.join(jv_clauses)} FROM judicial_violations"
            row = conn.execute(sql_jv, jv_params).fetchone()
            if row:
                for i, name in enumerate(adversary_names):
                    result[name]["violation_count"] = row[i] or 0

            # Impeachment counts per adversary
            imp_clauses = []
            imp_params = []
            for name in adversary_names:
                imp_clauses.append(
                    "SUM(CASE WHEN target LIKE ? OR evidence_summary LIKE ? THEN 1 ELSE 0 END)"
                )
                like = f"%{name.split()[0]}%"
                imp_params.extend([like, like])
            sql_imp = f"SELECT {', '.join(imp_clauses)} FROM impeachment_matrix"
            row = conn.execute(sql_imp, imp_params).fetchone()
            if row:
                for i, name in enumerate(adversary_names):
                    result[name]["impeachment_count"] = row[i] or 0

            # Top contradiction per adversary
            for name in adversary_names:
                like = f"%{name.split()[0]}%"
                ctr = conn.execute(
                    "SELECT contradiction_text, severity FROM contradiction_map "
                    "WHERE contradiction_text LIKE ? "
                    "ORDER BY severity DESC LIMIT 1",
                    (like,),
                ).fetchone()
                if ctr:
                    result[name]["top_contradiction"] = {
                        "text": ctr["contradiction_text"],
                        "severity": _safe_float(ctr["severity"]),
                    }

            return {"adversaries": result}
        except Exception as exc:
            return {"adversaries": {}, "error": str(exc)}

    def get_dashboard_data(self):
        """All analytics dashboard data in a single call.

        Consolidates separation counter, table counts, brain stats,
        top chains, and recent brain operations.
        """
        data = {
            "separation_days": _sep_days(),
            "separation_date": str(SEPARATION_DATE),
            "today": str(date.today()),
            "table_counts": {},
            "brain_stats": {},
            "top_chains": [],
            "recent_operations": [],
        }

        # Litigation DB table counts
        lit = self._lit()
        if lit:
            try:
                row = lit.execute(
                    "SELECT "
                    "(SELECT COUNT(*) FROM evidence_quotes) AS evidence_quotes, "
                    "(SELECT COUNT(*) FROM judicial_violations) AS judicial_violations, "
                    "(SELECT COUNT(*) FROM impeachment_matrix) AS impeachment_matrix, "
                    "(SELECT COUNT(*) FROM timeline_events) AS timeline_events, "
                    "(SELECT COUNT(*) FROM contradiction_map) AS contradiction_map"
                ).fetchone()
                data["table_counts"] = dict(row)
            except Exception as exc:
                data["table_counts"]["error"] = str(exc)

        # Brain DB stats
        brain = self._brain()
        if brain:
            try:
                row = brain.execute(
                    "SELECT "
                    "(SELECT COUNT(*) FROM nodes) AS nodes, "
                    "(SELECT COUNT(*) FROM edges) AS edges, "
                    "(SELECT COUNT(*) FROM chains) AS chains, "
                    "(SELECT COUNT(*) FROM gaps) AS gaps_total, "
                    "(SELECT COUNT(*) FROM gaps WHERE resolved = 0) AS gaps_open, "
                    "(SELECT COUNT(*) FROM gaps WHERE resolved = 1) AS gaps_resolved"
                ).fetchone()
                data["brain_stats"] = dict(row)
            except Exception as exc:
                data["brain_stats"]["error"] = str(exc)

            # Top 5 strongest chains
            try:
                top = brain.execute(
                    "SELECT id, chain_path, chain_type, strength_score, "
                    "lane, filing_id, length "
                    "FROM chains ORDER BY strength_score DESC LIMIT 5"
                ).fetchall()
                chains = []
                for r in top:
                    d = dict(r)
                    d["strength_score"] = _safe_float(d.get("strength_score"))
                    chains.append(d)
                data["top_chains"] = chains
            except Exception:
                pass

            # Recent brain operations (brain_log may be empty or absent)
            if _table_exists(brain, "brain_log"):
                try:
                    ops = brain.execute(
                        "SELECT * FROM brain_log ORDER BY rowid DESC LIMIT 10"
                    ).fetchall()
                    data["recent_operations"] = _rows_to_dicts(ops)
                except Exception:
                    pass

        return data

    def export_subgraph(self, node_ids):
        """Export a subgraph as D3.js-compatible JSON.

        Args:
            node_ids: List of node IDs to include.
        Returns:
            dict with ``nodes`` and ``links`` arrays matching D3 format.
        """
        conn = self._brain()
        if conn is None:
            return {"error": "Brain DB not available"}
        if not node_ids or not isinstance(node_ids, list):
            return {"error": "node_ids must be a non-empty list"}
        try:
            placeholders = ",".join("?" for _ in node_ids)

            # Fetch nodes
            node_rows = conn.execute(
                f"SELECT id, node_type, layer, label, description, "
                f"date_start, date_end, severity, confidence, readiness, "
                f"lane, metadata "
                f"FROM nodes WHERE id IN ({placeholders})",
                node_ids,
            ).fetchall()

            nodes = []
            id_set = set()
            for r in node_rows:
                d = dict(r)
                d["severity"] = _safe_float(d.get("severity"))
                d["confidence"] = _safe_float(d.get("confidence"))
                d["readiness"] = _safe_float(d.get("readiness"))
                # Parse metadata JSON if present
                raw_meta = d.get("metadata")
                if raw_meta:
                    try:
                        d["metadata"] = _json.loads(raw_meta)
                    except Exception:
                        pass
                nodes.append(d)
                id_set.add(d["id"])

            if not nodes:
                return {"nodes": [], "links": [], "warning": "No matching nodes found"}

            # Fetch edges between these nodes (both directions must be in set)
            edge_rows = conn.execute(
                f"SELECT id, source_id, target_id, edge_type, weight, evidence "
                f"FROM edges "
                f"WHERE source_id IN ({placeholders}) "
                f"AND target_id IN ({placeholders})",
                node_ids + node_ids,
            ).fetchall()

            links = []
            for r in edge_rows:
                d = dict(r)
                d["source"] = d.pop("source_id")
                d["target"] = d.pop("target_id")
                d["weight"] = _safe_float(d.get("weight"), 1.0)
                links.append(d)

            return {
                "nodes": nodes,
                "links": links,
                "node_count": len(nodes),
                "link_count": len(links),
            }
        except Exception as exc:
            return {"error": str(exc)}

    def search_everything(self, query, limit=50):
        """Unified search across both databases.

        Searches evidence_fts in litigation_context.db and nodes in
        mbp_brain.db, merges + deduplicates results, and sorts by
        relevance.

        Args:
            query: Search text.
            limit: Max results per source (total may be up to 2*limit
                   before dedup).
        Returns:
            dict with ``results`` list and ``count``.
        """
        if not query or not query.strip():
            return {"results": [], "count": 0, "method": "empty"}

        safe_q = _sanitize_fts(query)
        if not safe_q:
            return {"results": [], "count": 0, "method": "sanitized_empty"}

        all_results = []
        seen_ids = set()

        # --- litigation_context.db: evidence_fts ---
        lit = self._lit()
        if lit:
            # FTS5 path
            fts_hit = False
            if _table_exists(lit, "evidence_fts"):
                try:
                    rows = lit.execute(
                        "SELECT e.id, e.quote_text, e.source_file, e.category, "
                        "e.lane, e.relevance_score "
                        "FROM evidence_fts f "
                        "JOIN evidence_quotes e ON f.rowid = e.rowid "
                        "WHERE evidence_fts MATCH ? LIMIT ?",
                        (safe_q, limit),
                    ).fetchall()
                    for r in rows:
                        d = dict(r)
                        uid = f"lit_{d['id']}"
                        if uid not in seen_ids:
                            seen_ids.add(uid)
                            all_results.append({
                                "source": "litigation",
                                "id": d["id"],
                                "text": d.get("quote_text", ""),
                                "type": d.get("category", "evidence"),
                                "lane": d.get("lane", ""),
                                "score": _safe_float(d.get("relevance_score"), 0.5),
                                "source_file": d.get("source_file", ""),
                            })
                    fts_hit = len(rows) > 0
                except Exception:
                    pass

            # LIKE fallback if FTS5 missed
            if not fts_hit:
                like_pat = f"%{safe_q}%"
                try:
                    rows = lit.execute(
                        "SELECT id, quote_text, source_file, category, lane, "
                        "relevance_score FROM evidence_quotes "
                        "WHERE quote_text LIKE ? OR source_file LIKE ? "
                        "ORDER BY relevance_score DESC LIMIT ?",
                        (like_pat, like_pat, limit),
                    ).fetchall()
                    for r in rows:
                        d = dict(r)
                        uid = f"lit_{d['id']}"
                        if uid not in seen_ids:
                            seen_ids.add(uid)
                            all_results.append({
                                "source": "litigation",
                                "id": d["id"],
                                "text": d.get("quote_text", ""),
                                "type": d.get("category", "evidence"),
                                "lane": d.get("lane", ""),
                                "score": _safe_float(d.get("relevance_score"), 0.3),
                                "source_file": d.get("source_file", ""),
                            })
                except Exception:
                    pass

        # --- mbp_brain.db: nodes ---
        brain = self._brain()
        if brain:
            # FTS5 path
            brain_fts_hit = False
            if _table_exists(brain, "nodes_fts"):
                try:
                    rows = brain.execute(
                        "SELECT n.id, n.label, n.node_type, n.layer, n.lane, "
                        "n.confidence "
                        "FROM nodes_fts f "
                        "JOIN nodes n ON f.rowid = n.rowid "
                        "WHERE nodes_fts MATCH ? LIMIT ?",
                        (safe_q, limit),
                    ).fetchall()
                    for r in rows:
                        d = dict(r)
                        uid = f"brain_{d['id']}"
                        if uid not in seen_ids:
                            seen_ids.add(uid)
                            all_results.append({
                                "source": "brain",
                                "id": d["id"],
                                "text": d.get("label", ""),
                                "type": d.get("node_type", ""),
                                "lane": d.get("lane", ""),
                                "score": _safe_float(d.get("confidence"), 0.5),
                                "layer": d.get("layer", ""),
                            })
                    brain_fts_hit = len(rows) > 0
                except Exception:
                    pass

            # LIKE fallback
            if not brain_fts_hit:
                like_pat = f"%{safe_q}%"
                try:
                    rows = brain.execute(
                        "SELECT id, label, node_type, layer, lane, confidence "
                        "FROM nodes WHERE label LIKE ? OR description LIKE ? "
                        "ORDER BY confidence DESC LIMIT ?",
                        (like_pat, like_pat, limit),
                    ).fetchall()
                    for r in rows:
                        d = dict(r)
                        uid = f"brain_{d['id']}"
                        if uid not in seen_ids:
                            seen_ids.add(uid)
                            all_results.append({
                                "source": "brain",
                                "id": d["id"],
                                "text": d.get("label", ""),
                                "type": d.get("node_type", ""),
                                "lane": d.get("lane", ""),
                                "score": _safe_float(d.get("confidence"), 0.3),
                                "layer": d.get("layer", ""),
                            })
                except Exception:
                    pass

        # Sort by score descending, then trim
        all_results.sort(key=lambda x: x.get("score", 0), reverse=True)
        all_results = all_results[: limit * 2]

        return {
            "results": all_results,
            "count": len(all_results),
        }

    # ===================================================================
    # T1: BRAIN REBUILD (build_mbp_brain.py)
    # ===================================================================

    def rebuild_brain(self):
        """Trigger a full brain rebuild (T1 -- build_mbp_brain.py).

        Runs asynchronously in a background thread.  Returns immediately
        with status ``started``.  Poll ``get_health()`` to see when
        new node/edge counts appear.
        """
        script = REPO_ROOT / "scripts" / "build_mbp_brain.py"
        if not script.exists():
            return {"status": "error", "message": f"Script not found: {script}"}

        def _run():
            try:
                subprocess.run(
                    [_PYTHON, "-I", str(script)],
                    cwd=str(REPO_ROOT),
                    timeout=600,
                    capture_output=True,
                )
            except Exception:
                pass

        t = threading.Thread(target=_run, daemon=True, name="brain-rebuild")
        t.start()
        return {"status": "started", "script": str(script)}

    # ===================================================================
    # T2: CHAIN RECOMPUTE (compute_chains.py)
    # ===================================================================

    def recompute_chains(self):
        """Trigger chain recomputation (T2 -- compute_chains.py)."""
        script = REPO_ROOT / "scripts" / "compute_chains.py"
        if not script.exists():
            return {"status": "error", "message": f"Script not found: {script}"}

        def _run():
            try:
                subprocess.run(
                    [_PYTHON, "-I", str(script)],
                    cwd=str(REPO_ROOT),
                    timeout=600,
                    capture_output=True,
                )
            except Exception:
                pass

        t = threading.Thread(target=_run, daemon=True, name="chain-recompute")
        t.start()
        return {"status": "started", "script": str(script)}

    # ===================================================================
    # T3: FILE WATCHER (brain_watcher.py)
    # ===================================================================

    def start_watcher(self):
        """Start the MEEK file watcher daemon (T3 -- brain_watcher.py).

        Launches brain_watcher.py as a background subprocess with
        auto-routing enabled.  Returns immediately.
        """
        script = REPO_ROOT / "scripts" / "brain_watcher.py"
        if not script.exists():
            return {"status": "error", "message": f"Script not found: {script}"}

        def _run():
            try:
                subprocess.Popen(
                    [_PYTHON, "-I", str(script)],
                    cwd=str(REPO_ROOT),
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                )
            except Exception:
                pass

        t = threading.Thread(target=_run, daemon=True, name="watcher-start")
        t.start()
        return {"status": "started", "script": str(script)}

    # ===================================================================
    # CLOSED LOOP ORCHESTRATOR
    # ===================================================================

    def run_full_cycle(self):
        """Execute the complete closed-loop pipeline:

        Evidence -> Violation -> Authority -> Remedy -> Filing -> back to Evidence

        Runs T1 (brain build) -> T2 (chains) -> T5 (evolution) -> T6 (export)
        sequentially in a background thread.  Returns immediately.
        """
        def _cycle():
            scripts = [
                ("T1-Brain", REPO_ROOT / "scripts" / "build_mbp_brain.py"),
                ("T2-Chains", REPO_ROOT / "scripts" / "compute_chains.py"),
                ("T5-Evolution", REPO_ROOT / "scripts" / "brain_evolution.py"),
                ("T6-Export", REPO_ROOT / "scripts" / "export_brain_d3.py"),
            ]
            for label, script in scripts:
                if script.exists():
                    try:
                        subprocess.run(
                            [_PYTHON, "-I", str(script)],
                            cwd=str(REPO_ROOT),
                            timeout=600,
                            capture_output=True,
                        )
                    except Exception:
                        pass

        t = threading.Thread(target=_cycle, daemon=True, name="full-cycle")
        t.start()
        return {
            "status": "started",
            "pipeline": "T1-Brain -> T2-Chains -> T5-Evolution -> T6-Export",
        }

    # ===================================================================
    # COMMUNITY INTELLIGENCE (v15.0 — Leiden Clusters)
    # ===================================================================

    def get_community_intel(self, community_id):
        """Full intelligence for a community: metadata, top nodes, patterns."""
        conn = self._brain()
        if conn is None:
            return {"error": "Brain DB not available"}
        try:
            comm = conn.execute(
                "SELECT * FROM communities WHERE id = ?", (community_id,)
            ).fetchone()
            if not comm:
                return {"error": f"Community {community_id} not found"}

            members = conn.execute(
                "SELECT n.id, n.label, n.node_type, n.lane, na.pagerank, "
                "na.hub_score, na.authority_score "
                "FROM community_members cm "
                "JOIN nodes n ON cm.node_id = n.id "
                "LEFT JOIN node_analytics na ON n.id = na.node_id "
                "WHERE cm.community_id = ? "
                "ORDER BY COALESCE(na.pagerank, 0) DESC LIMIT 50",
                (community_id,),
            ).fetchall()

            children = conn.execute(
                "SELECT id, label, member_count, level FROM communities "
                "WHERE parent_id = ? ORDER BY member_count DESC",
                (community_id,),
            ).fetchall()

            edges = conn.execute(
                "SELECT ce.*, c2.label AS target_label "
                "FROM community_edges ce "
                "LEFT JOIN communities c2 ON ce.target_id = c2.id "
                "WHERE ce.source_id = ? OR ce.target_id = ? "
                "ORDER BY ce.total_weight DESC LIMIT 30",
                (community_id, community_id),
            ).fetchall()

            return {
                "community": dict(comm),
                "top_nodes": _rows_to_dicts(members),
                "children": _rows_to_dicts(children),
                "inter_edges": _rows_to_dicts(edges),
                "node_count": len(members),
            }
        except Exception as exc:
            return {"error": str(exc)}

    def search_communities(self, query, limit=30):
        """Search communities by label, narrative, or key actors."""
        conn = self._brain()
        if conn is None:
            return {"error": "Brain DB not available"}
        try:
            safe_q = _sanitize_fts(query)
            like_pat = f"%{safe_q}%"
            rows = conn.execute(
                "SELECT id, label, level, lane, member_count, "
                "evidence_strength, authority_completeness, impeachment_density "
                "FROM communities "
                "WHERE label LIKE ? OR narrative LIKE ? OR key_actors LIKE ? "
                "ORDER BY member_count DESC LIMIT ?",
                (like_pat, like_pat, like_pat, limit),
            ).fetchall()
            return {"communities": _rows_to_dicts(rows), "count": len(rows)}
        except Exception as exc:
            return {"error": str(exc)}

    def get_community_timeline(self, community_id, limit=100):
        """Chronological events within a community."""
        conn = self._brain()
        if conn is None:
            return {"error": "Brain DB not available"}
        try:
            rows = conn.execute(
                "SELECT n.id, n.label, n.node_type, n.date_start, n.lane "
                "FROM community_members cm "
                "JOIN nodes n ON cm.node_id = n.id "
                "WHERE cm.community_id = ? AND n.date_start IS NOT NULL "
                "ORDER BY n.date_start ASC LIMIT ?",
                (community_id, limit),
            ).fetchall()
            return {"events": _rows_to_dicts(rows), "count": len(rows)}
        except Exception as exc:
            return {"error": str(exc)}

    def expand_community(self, community_id, limit=200):
        """Return individual nodes for JS-side expansion of a community."""
        conn = self._brain()
        if conn is None:
            return {"error": "Brain DB not available"}
        try:
            nodes = conn.execute(
                "SELECT n.id, n.label, n.node_type, n.lane, n.date_start, "
                "na.pagerank, na.hub_score, na.authority_score "
                "FROM community_members cm "
                "JOIN nodes n ON cm.node_id = n.id "
                "LEFT JOIN node_analytics na ON n.id = na.node_id "
                "WHERE cm.community_id = ? "
                "ORDER BY COALESCE(na.pagerank, 0) DESC LIMIT ?",
                (community_id, limit),
            ).fetchall()

            node_ids = [r["id"] for r in nodes]
            edges = []
            if node_ids:
                placeholders = ",".join("?" * len(node_ids))
                edges = conn.execute(
                    f"SELECT source_id, target_id, edge_type, weight "
                    f"FROM edges "
                    f"WHERE source_id IN ({placeholders}) "
                    f"AND target_id IN ({placeholders}) LIMIT 1000",
                    node_ids + node_ids,
                ).fetchall()

            return {
                "nodes": _rows_to_dicts(nodes),
                "edges": _rows_to_dicts(edges),
                "community_id": community_id,
            }
        except Exception as exc:
            return {"error": str(exc)}

    def get_community_stats(self):
        """Summary stats for the community system."""
        conn = self._brain()
        if conn is None:
            return {"error": "Brain DB not available"}
        try:
            row = conn.execute(
                "SELECT "
                "(SELECT COUNT(*) FROM communities WHERE level = 0) AS lanes, "
                "(SELECT COUNT(*) FROM communities WHERE level = 1) AS epochs, "
                "(SELECT COUNT(*) FROM communities WHERE level = 2) AS communities, "
                "(SELECT COUNT(*) FROM community_members) AS memberships, "
                "(SELECT COUNT(*) FROM community_edges) AS inter_edges, "
                "(SELECT COUNT(*) FROM node_analytics WHERE pagerank > 0) AS ranked_nodes"
            ).fetchone()
            return dict(row) if row else {}
        except Exception as exc:
            return {"error": str(exc)}

    def get_detected_patterns(self, limit=50):
        """Autonomously detected patterns across communities."""
        conn = self._brain()
        if conn is None:
            return {"error": "Brain DB not available"}
        try:
            if not _table_exists(conn, "detected_patterns"):
                return {"patterns": [], "count": 0}
            rows = conn.execute(
                "SELECT * FROM detected_patterns "
                "ORDER BY confidence DESC, detected_at DESC LIMIT ?",
                (limit,),
            ).fetchall()
            return {"patterns": _rows_to_dicts(rows), "count": len(rows)}
        except Exception as exc:
            return {"error": str(exc)}

    # ===================================================================
    # PATTERN & EMBEDDING INTELLIGENCE (v15.0 Wave 3)
    # ===================================================================

    def get_pattern_summary(self):
        """Aggregate pattern stats: counts by type and severity."""
        conn = self._brain()
        if conn is None:
            return {"error": "Brain DB not available"}
        try:
            if not _table_exists(conn, "detected_patterns"):
                return {"summary": [], "total": 0}
            rows = conn.execute(
                "SELECT pattern_type, severity, COUNT(*) AS cnt, "
                "ROUND(AVG(confidence), 2) AS avg_confidence "
                "FROM detected_patterns "
                "GROUP BY pattern_type, severity "
                "ORDER BY cnt DESC"
            ).fetchall()
            total = conn.execute(
                "SELECT COUNT(*) FROM detected_patterns"
            ).fetchone()[0]
            return {"summary": _rows_to_dicts(rows), "total": total}
        except Exception as exc:
            return {"error": str(exc)}

    def get_patterns_by_type(self, pattern_type, severity=None, limit=50):
        """Get patterns filtered by type and optional severity."""
        conn = self._brain()
        if conn is None:
            return {"error": "Brain DB not available"}
        try:
            if not _table_exists(conn, "detected_patterns"):
                return {"patterns": [], "count": 0}
            sql = "SELECT * FROM detected_patterns WHERE pattern_type = ?"
            params = [pattern_type]
            if severity:
                sql += " AND severity = ?"
                params.append(severity)
            sql += " ORDER BY confidence DESC LIMIT ?"
            params.append(limit)
            rows = conn.execute(sql, params).fetchall()
            return {"patterns": _rows_to_dicts(rows), "count": len(rows)}
        except Exception as exc:
            return {"error": str(exc)}

    def get_cross_lane_bridges(self, limit=50):
        """Get nodes or communities that bridge multiple litigation lanes."""
        conn = self._brain()
        if conn is None:
            return {"error": "Brain DB not available"}
        try:
            # Strategy: find nodes that appear in communities across 2+ lanes
            rows = conn.execute(
                "SELECT n.id, n.label, n.node_type, "
                "GROUP_CONCAT(DISTINCT c.lane) AS lanes, "
                "COUNT(DISTINCT c.lane) AS lane_count, "
                "GROUP_CONCAT(DISTINCT c.id) AS community_ids "
                "FROM community_members cm "
                "JOIN nodes n ON cm.node_id = n.id "
                "JOIN communities c ON cm.community_id = c.id "
                "WHERE c.level = 2 AND c.lane IS NOT NULL "
                "GROUP BY n.id "
                "HAVING COUNT(DISTINCT c.lane) >= 2 "
                "ORDER BY lane_count DESC, n.label "
                "LIMIT ?",
                (limit,),
            ).fetchall()
            # Also check detected_patterns for CROSS_LANE_BRIDGE type
            bridge_patterns = []
            if _table_exists(conn, "detected_patterns"):
                bridge_patterns = _rows_to_dicts(
                    conn.execute(
                        "SELECT * FROM detected_patterns "
                        "WHERE pattern_type = 'CROSS_LANE_BRIDGE' "
                        "ORDER BY confidence DESC LIMIT ?",
                        (limit,),
                    ).fetchall()
                )
            return {
                "bridge_nodes": _rows_to_dicts(rows),
                "bridge_patterns": bridge_patterns,
                "count": len(rows),
            }
        except Exception as exc:
            return {"error": str(exc)}

    def search_similar_communities(self, query, top_k=10):
        """Semantic vector search: find communities similar to a query string."""
        try:
            import lancedb
            from sentence_transformers import SentenceTransformer

            lance_dir = REPO_ROOT / "00_SYSTEM" / "engines" / "semantic" / "data"
            db = lancedb.connect(str(lance_dir))
            table_names = db.table_names()
            if "sem_communities" not in table_names:
                return {"error": "sem_communities table not found in LanceDB"}

            tbl = db.open_table("sem_communities")
            model = SentenceTransformer("all-MiniLM-L6-v2")
            q_vec = model.encode(query).tolist()
            results = tbl.search(q_vec).limit(top_k).to_pandas()

            communities = []
            for _, row in results.iterrows():
                communities.append({
                    "community_id": row.get("community_id", ""),
                    "label": row.get("label", ""),
                    "lane": row.get("lane", ""),
                    "member_count": int(row.get("member_count", 0)),
                    "text": row.get("text", "")[:300],
                    "distance": float(row.get("_distance", 0)),
                })
            return {"communities": communities, "query": query, "count": len(communities)}
        except ImportError:
            # Fallback: keyword search if LanceDB/sentence-transformers not available
            return self.search_communities(query, limit=top_k)
        except Exception as exc:
            return {"error": str(exc)}

    def build_filing_section(self, community_id):
        """Auto-generate a filing-ready evidence section from community data."""
        conn = self._brain()
        if conn is None:
            return {"error": "Brain DB not available"}
        try:
            comm = conn.execute(
                "SELECT * FROM communities WHERE id = ?", (community_id,)
            ).fetchone()
            if not comm:
                return {"error": f"Community {community_id} not found"}
            comm_d = dict(comm)

            # Get top evidence nodes (quotes, facts, events)
            evidence_nodes = conn.execute(
                "SELECT n.id, n.label, n.node_type, n.description, "
                "n.date_start, n.severity, n.confidence, n.lane, "
                "na.pagerank "
                "FROM community_members cm "
                "JOIN nodes n ON cm.node_id = n.id "
                "LEFT JOIN node_analytics na ON n.id = na.node_id "
                "WHERE cm.community_id = ? "
                "AND n.node_type IN ('Quote', 'Fact', 'Event', 'Contradiction') "
                "ORDER BY COALESCE(na.pagerank, 0) DESC LIMIT 25",
                (community_id,),
            ).fetchall()

            # Get authority nodes
            authority_nodes = conn.execute(
                "SELECT n.id, n.label, n.node_type, n.description "
                "FROM community_members cm "
                "JOIN nodes n ON cm.node_id = n.id "
                "WHERE cm.community_id = ? "
                "AND n.node_type IN ('Statute', 'CourtRule', 'CaseLaw') "
                "ORDER BY n.label LIMIT 15",
                (community_id,),
            ).fetchall()

            # Get related patterns
            patterns = []
            if _table_exists(conn, "detected_patterns"):
                patterns = _rows_to_dicts(
                    conn.execute(
                        "SELECT pattern_type, description, severity, confidence "
                        "FROM detected_patterns "
                        "WHERE communities_involved LIKE ? "
                        "ORDER BY confidence DESC LIMIT 10",
                        (f"%{community_id}%",),
                    ).fetchall()
                )

            # Build narrative text
            lane = comm_d.get("lane", "Unknown")
            label = comm_d.get("label", "Unknown Community")
            date_range = ""
            if comm_d.get("date_start") and comm_d.get("date_end"):
                date_range = f" ({comm_d['date_start']} to {comm_d['date_end']})"

            section_lines = [
                f"## {label}{date_range}",
                f"**Lane**: {lane} | **Members**: {comm_d.get('member_count', 0)}",
                "",
            ]

            if evidence_nodes:
                section_lines.append("### Evidence")
                for n in evidence_nodes:
                    nd = dict(n)
                    date_str = f" [{nd.get('date_start', '')}]" if nd.get("date_start") else ""
                    desc = nd.get("description", nd.get("label", ""))
                    if desc and len(desc) > 200:
                        desc = desc[:200] + "..."
                    section_lines.append(
                        f"- **{nd.get('node_type', '')}**{date_str}: {desc}"
                    )

            if authority_nodes:
                section_lines.append("")
                section_lines.append("### Legal Authorities")
                for n in authority_nodes:
                    nd = dict(n)
                    section_lines.append(f"- {nd.get('label', '')} ({nd.get('node_type', '')})")

            if patterns:
                section_lines.append("")
                section_lines.append("### Detected Patterns")
                for p in patterns:
                    section_lines.append(
                        f"- **{p['pattern_type']}** [{p['severity']}]: {p['description'][:150]}"
                    )

            return {
                "community": comm_d,
                "section_markdown": "\n".join(section_lines),
                "evidence_count": len(evidence_nodes),
                "authority_count": len(authority_nodes),
                "pattern_count": len(patterns),
            }
        except Exception as exc:
            return {"error": str(exc)}

    # ===================================================================
    # HEALTH CHECK
    # ===================================================================

    def get_health(self):
        """Brain health check with comprehensive counts."""
        brain = self._brain()
        lit = self._lit()
        health = {
            "version": VERSION,
            "separation_days": _sep_days(),
            "brain_db": {"available": brain is not None, "path": str(BRAIN_DB)},
            "lit_db": {"available": lit is not None, "path": str(LIT_DB)},
        }

        if brain:
            try:
                row = brain.execute(
                    "SELECT "
                    "(SELECT COUNT(*) FROM nodes) AS nodes, "
                    "(SELECT COUNT(*) FROM edges) AS edges, "
                    "(SELECT COUNT(*) FROM chains) AS chains, "
                    "(SELECT COUNT(*) FROM gaps WHERE resolved = 0) AS open_gaps, "
                    "(SELECT MAX(version) FROM versions) AS version"
                ).fetchone()
                health["brain_db"].update(dict(row))

                # Community stats (v15.0)
                if _table_exists(brain, "communities"):
                    crow = brain.execute(
                        "SELECT "
                        "(SELECT COUNT(*) FROM communities) AS communities, "
                        "(SELECT COUNT(*) FROM community_members) AS memberships, "
                        "(SELECT COUNT(*) FROM node_analytics WHERE pagerank > 0) AS ranked_nodes, "
                        + ("(SELECT COUNT(*) FROM detected_patterns) AS patterns, "
                           if _table_exists(brain, "detected_patterns") else "0 AS patterns, ")
                        + ("(SELECT COUNT(*) FROM community_embeddings) AS embeddings"
                           if _table_exists(brain, "community_embeddings") else "0 AS embeddings")
                    ).fetchone()
                    health["brain_db"]["communities"] = dict(crow) if crow else {}
            except Exception as exc:
                health["brain_db"]["error"] = str(exc)

        if lit:
            try:
                row = lit.execute(
                    "SELECT "
                    "(SELECT COUNT(*) FROM evidence_quotes) AS evidence, "
                    "(SELECT COUNT(*) FROM judicial_violations) AS violations, "
                    "(SELECT COUNT(*) FROM timeline_events) AS timeline, "
                    "(SELECT COUNT(*) FROM impeachment_matrix) AS impeachment"
                ).fetchone()
                health["lit_db"].update(dict(row))
            except Exception as exc:
                health["lit_db"]["error"] = str(exc)

        health["graph_json"] = {
            "exists": GRAPH_JSON.exists(),
            "size_mb": round(GRAPH_JSON.stat().st_size / 1048576, 1) if GRAPH_JSON.exists() else 0,
        }
        return health

    # ===================================================================
    # ENGINE STATUS — which satellite engines loaded successfully
    # ===================================================================

    def get_engine_status(self):
        engines = {
            "athena": {"available": HAS_ATHENA, "loaded": self._athena is not None,
                       "methods": ["doctrine_map","enrich_argument","find_authority","citation_network",
                                   "scotus_precedent","michigan_precedent","normalize_language",
                                   "authority_score","build_authority_brief","professional_standards"]},
            "chronos": {"available": HAS_CHRONOS, "loaded": self._chronos is not None,
                        "methods": ["build_timeline","detect_cycles","escalation_score","separation_impact",
                                    "event_velocity","temporal_clusters","milestone_timeline",
                                    "actor_timeline","lane_chronology"]},
            "cortex": {"available": HAS_CORTEX, "loaded": self._cortex is not None,
                       "methods": ["irac_analyze","compute_damages","filing_qa","build_narrative",
                                   "build_xexam","trace_authority","lane_readiness","adversary_profile",
                                   "check_deadlines","detect_gaps","generate_rebuttal","nexus_fuse",
                                   "strategic_report","search_all"]},
            "oracle": {"available": HAS_ORACLE, "loaded": self._oracle is not None,
                       "methods": ["predict_ruling","forecast_adversary","filing_success_probability",
                                   "optimal_sequence","counter_strategy","risk_matrix","early_warning",
                                   "judicial_tendency"]},
            "prometheus": {"available": HAS_PROMETHEUS, "loaded": self._prometheus is not None,
                           "methods": ["generate_irac","authority_chain","build_affidavit","exhibit_index",
                                       "strike_plan","weapon_inventory","xexam_script",
                                       "filing_package_status","rebuttal_builder"]},
            "automaton": {"available": HAS_AUTOMATON, "loaded": self._automaton is not None,
                          "methods": ["start","stop","status","get_results","get_templates","get_centrality"]},
            "backend": {"available": HAS_BACKEND, "loaded": self._backend is not None,
                        "methods": ["capabilities","fraud_detect","draft_complaint","draft_injunction",
                                    "draft_show_cause","draft_discovery","draft_disqualification",
                                    "draft_jtc","draft_appeal","draft_superintending","draft_relief",
                                    "bates_manifest","rules_tactical","pdf_stamp","pdf_normalize",
                                    "forms_resolve","forms_catalog"]},
            "bleeding_edge": {"available": HAS_BLEEDING_EDGE, "loaded": self._bleeding_edge is not None,
                              "methods": ["capabilities","evidence_heatmap","authority_coverage",
                                          "impeachment_arsenal","filing_readiness","judicial_dashboard",
                                          "contradiction_network","separation_counter","full_dashboard",
                                          "semantic_search","semantic_encode","lance_stats",
                                          "hybrid_search","get_palette","list_palettes",
                                          "scan_database","merge_databases","export_graph_json",
                                          "export_graph_csv","build_csv_graph",
                                          "hash_file","scan_directory_changes"]},
            "intelligence": {"available": HAS_INTELLIGENCE, "loaded": self._intelligence is not None,
                             "engines": ["nexus","nemesis","cerberus","chimera","temporal",
                                         "causal","damages","irac","hypergraph","narrative","forge"],
                             "methods": ["nexus_fuse","nexus_argue","nexus_readiness","nexus_damages",
                                         "nexus_red_team","nexus_case_map","nexus_priorities",
                                         "nemesis_profile","nemesis_predict","nemesis_vulnerabilities",
                                         "cerberus_report","cerberus_gaps","cerberus_weapons",
                                         "chimera_detect","chimera_impeachment","chimera_patterns",
                                         "temporal_poisonous_tree","temporal_conspiracy","temporal_anomalies",
                                         "causal_poisonous_tree","causal_retaliation","causal_for_filing",
                                         "damages_total","damages_by_lane","damages_summary",
                                         "irac_argument","irac_strong","irac_gaps",
                                         "hypergraph_conspiracy","hypergraph_strongest","hypergraph_filing_ammo",
                                         "narrative_claim","narrative_defendant","narrative_exhibits",
                                         "forge_assemble","forge_validate","forge_list_packets"]},
            "operations": {"available": HAS_OPERATIONS, "loaded": self._operations is not None,
                           "engines": ["orchestrator","qa","filing_engine","filing_assembler",
                                       "analyzer","rebuttal","perception","apotheosis","lexicon"],
                           "methods": ["orchestrator_fuse","orchestrator_entity_profile","orchestrator_filing_intel",
                                       "orchestrator_health","qa_audit_all","qa_summary",
                                       "filing_scan_triggers","filing_run","filing_engine_status",
                                       "assembler_build_stack","analyze_document",
                                       "rebuttal_all","rebuttal_by_factor","rebuttal_appellate_section",
                                       "perception_classify","perception_ner","perception_embed",
                                       "apotheosis_health","apotheosis_gaps","apotheosis_recommendations",
                                       "apotheosis_trajectory","lexicon_filing_roadmap","lexicon_compute_deadlines",
                                       "lexicon_evidence_check","lexicon_canon_analysis","lexicon_ask"]},
        }
        engines["_summary"] = {
            "total": 10,
            "available": sum(1 for e in engines.values() if isinstance(e, dict) and e.get("available")),
            "loaded": sum(1 for e in engines.values() if isinstance(e, dict) and e.get("loaded")),
            "version": VERSION,
        }
        return engines

    # ===================================================================
    # ATHENA API — Legal Authority & Doctrine Intelligence
    # ===================================================================

    def athena_doctrine_map(self, doctrine=None):
        eng = self._get_athena()
        if not eng:
            return {"error": "Athena engine not available"}
        try:
            return eng.doctrine_map(doctrine)
        except Exception as e:
            return {"error": str(e)}

    def athena_enrich_argument(self, argument, lane=None):
        eng = self._get_athena()
        if not eng:
            return {"error": "Athena engine not available"}
        try:
            return eng.enrich_argument(argument, lane=lane)
        except Exception as e:
            return {"error": str(e)}

    def athena_find_authority(self, topic, limit=10):
        eng = self._get_athena()
        if not eng:
            return {"error": "Athena engine not available"}
        try:
            return eng.find_authority(topic, limit=limit)
        except Exception as e:
            return {"error": str(e)}

    def athena_citation_network(self, citation=None):
        eng = self._get_athena()
        if not eng:
            return {"error": "Athena engine not available"}
        try:
            return eng.citation_network(citation)
        except Exception as e:
            return {"error": str(e)}

    def athena_scotus_precedent(self, topic):
        eng = self._get_athena()
        if not eng:
            return {"error": "Athena engine not available"}
        try:
            return eng.scotus_precedent(topic)
        except Exception as e:
            return {"error": str(e)}

    def athena_michigan_precedent(self, topic):
        eng = self._get_athena()
        if not eng:
            return {"error": "Athena engine not available"}
        try:
            return eng.michigan_precedent(topic)
        except Exception as e:
            return {"error": str(e)}

    def athena_normalize_language(self, text):
        eng = self._get_athena()
        if not eng:
            return {"error": "Athena engine not available"}
        try:
            return eng.normalize_language(text)
        except Exception as e:
            return {"error": str(e)}

    def athena_authority_score(self, citation):
        eng = self._get_athena()
        if not eng:
            return {"error": "Athena engine not available"}
        try:
            return eng.authority_score(citation)
        except Exception as e:
            return {"error": str(e)}

    def athena_build_authority_brief(self, topic, lane=None):
        eng = self._get_athena()
        if not eng:
            return {"error": "Athena engine not available"}
        try:
            return eng.build_authority_brief(topic, lane=lane)
        except Exception as e:
            return {"error": str(e)}

    def athena_professional_standards(self, category=None):
        eng = self._get_athena()
        if not eng:
            return {"error": "Athena engine not available"}
        try:
            return eng.professional_standards(category)
        except Exception as e:
            return {"error": str(e)}

    # ===================================================================
    # CHRONOS API — Temporal Pattern Mining
    # ===================================================================

    def chronos_build_timeline(self, lane=None, limit=100):
        eng = self._get_chronos()
        if not eng:
            return {"error": "Chronos engine not available"}
        try:
            return eng.build_timeline(lane=lane, limit=limit)
        except Exception as e:
            return {"error": str(e)}

    def chronos_detect_cycles(self, lane=None):
        eng = self._get_chronos()
        if not eng:
            return {"error": "Chronos engine not available"}
        try:
            return eng.detect_cycles(lane=lane)
        except Exception as e:
            return {"error": str(e)}

    def chronos_escalation_score(self, actor=None):
        eng = self._get_chronos()
        if not eng:
            return {"error": "Chronos engine not available"}
        try:
            return eng.escalation_score(actor=actor)
        except Exception as e:
            return {"error": str(e)}

    def chronos_separation_impact(self):
        eng = self._get_chronos()
        if not eng:
            return {"error": "Chronos engine not available"}
        try:
            return eng.separation_impact()
        except Exception as e:
            return {"error": str(e)}

    def chronos_event_velocity(self, lane=None, window_days=30):
        eng = self._get_chronos()
        if not eng:
            return {"error": "Chronos engine not available"}
        try:
            return eng.event_velocity(lane=lane, window_days=window_days)
        except Exception as e:
            return {"error": str(e)}

    def chronos_temporal_clusters(self, n_clusters=5):
        eng = self._get_chronos()
        if not eng:
            return {"error": "Chronos engine not available"}
        try:
            return eng.temporal_clusters(n_clusters=n_clusters)
        except Exception as e:
            return {"error": str(e)}

    def chronos_milestone_timeline(self):
        eng = self._get_chronos()
        if not eng:
            return {"error": "Chronos engine not available"}
        try:
            return eng.milestone_timeline()
        except Exception as e:
            return {"error": str(e)}

    def chronos_actor_timeline(self, actor):
        eng = self._get_chronos()
        if not eng:
            return {"error": "Chronos engine not available"}
        try:
            return eng.actor_timeline(actor)
        except Exception as e:
            return {"error": str(e)}

    def chronos_lane_chronology(self, lane):
        eng = self._get_chronos()
        if not eng:
            return {"error": "Chronos engine not available"}
        try:
            return eng.lane_chronology(lane)
        except Exception as e:
            return {"error": str(e)}

    # ===================================================================
    # ORACLE API — Predictive Intelligence
    # ===================================================================

    def oracle_predict_ruling(self, motion_type, lane=None):
        eng = self._get_oracle()
        if not eng:
            return {"error": "Oracle engine not available"}
        try:
            return eng.predict_ruling(motion_type, lane=lane)
        except Exception as e:
            return {"error": str(e)}

    def oracle_forecast_adversary(self, adversary=None):
        eng = self._get_oracle()
        if not eng:
            return {"error": "Oracle engine not available"}
        try:
            return eng.forecast_adversary(adversary=adversary)
        except Exception as e:
            return {"error": str(e)}

    def oracle_filing_success(self, filing_id):
        eng = self._get_oracle()
        if not eng:
            return {"error": "Oracle engine not available"}
        try:
            return eng.filing_success_probability(filing_id)
        except Exception as e:
            return {"error": str(e)}

    def oracle_optimal_sequence(self):
        eng = self._get_oracle()
        if not eng:
            return {"error": "Oracle engine not available"}
        try:
            return eng.optimal_sequence()
        except Exception as e:
            return {"error": str(e)}

    def oracle_counter_strategy(self, adversary_action):
        eng = self._get_oracle()
        if not eng:
            return {"error": "Oracle engine not available"}
        try:
            return eng.counter_strategy(adversary_action)
        except Exception as e:
            return {"error": str(e)}

    def oracle_risk_matrix(self, lane=None):
        eng = self._get_oracle()
        if not eng:
            return {"error": "Oracle engine not available"}
        try:
            return eng.risk_matrix(lane=lane)
        except Exception as e:
            return {"error": str(e)}

    def oracle_early_warning(self):
        eng = self._get_oracle()
        if not eng:
            return {"error": "Oracle engine not available"}
        try:
            return eng.early_warning()
        except Exception as e:
            return {"error": str(e)}

    def oracle_judicial_tendency(self, judge=None):
        eng = self._get_oracle()
        if not eng:
            return {"error": "Oracle engine not available"}
        try:
            return eng.judicial_tendency(judge=judge)
        except Exception as e:
            return {"error": str(e)}

    # ===================================================================
    # PROMETHEUS API — Strategic Weapon Delivery
    # ===================================================================

    def prometheus_generate_irac(self, issue, lane=None):
        eng = self._get_prometheus()
        if not eng:
            return {"error": "Prometheus engine not available"}
        try:
            return eng.generate_irac(issue, lane=lane)
        except Exception as e:
            return {"error": str(e)}

    def prometheus_authority_chain(self, topic):
        eng = self._get_prometheus()
        if not eng:
            return {"error": "Prometheus engine not available"}
        try:
            return eng.authority_chain(topic)
        except Exception as e:
            return {"error": str(e)}

    def prometheus_build_affidavit(self, claims, lane=None):
        eng = self._get_prometheus()
        if not eng:
            return {"error": "Prometheus engine not available"}
        try:
            return eng.build_affidavit(claims, lane=lane)
        except Exception as e:
            return {"error": str(e)}

    def prometheus_exhibit_index(self, lane=None):
        eng = self._get_prometheus()
        if not eng:
            return {"error": "Prometheus engine not available"}
        try:
            return eng.exhibit_index(lane=lane)
        except Exception as e:
            return {"error": str(e)}

    def prometheus_strike_plan(self, target=None):
        eng = self._get_prometheus()
        if not eng:
            return {"error": "Prometheus engine not available"}
        try:
            return eng.strike_plan(target=target)
        except Exception as e:
            return {"error": str(e)}

    def prometheus_weapon_inventory(self, lane=None):
        eng = self._get_prometheus()
        if not eng:
            return {"error": "Prometheus engine not available"}
        try:
            return eng.weapon_inventory(lane=lane)
        except Exception as e:
            return {"error": str(e)}

    def prometheus_xexam_script(self, witness, topic=None):
        eng = self._get_prometheus()
        if not eng:
            return {"error": "Prometheus engine not available"}
        try:
            return eng.xexam_script(witness, topic=topic)
        except Exception as e:
            return {"error": str(e)}

    def prometheus_filing_package_status(self, lane=None):
        eng = self._get_prometheus()
        if not eng:
            return {"error": "Prometheus engine not available"}
        try:
            return eng.filing_package_status(lane=lane)
        except Exception as e:
            return {"error": str(e)}

    def prometheus_rebuttal_builder(self, claim):
        eng = self._get_prometheus()
        if not eng:
            return {"error": "Prometheus engine not available"}
        try:
            return eng.rebuttal_builder(claim)
        except Exception as e:
            return {"error": str(e)}

    # ===================================================================
    # CORTEX API — 27-Engine Orchestrator
    # ===================================================================

    def cortex_irac_analyze(self, issue, lane=None):
        eng = self._get_cortex()
        if not eng:
            return {"error": "Cortex engine not available"}
        try:
            return eng.irac_analyze(issue, lane=lane)
        except Exception as e:
            return {"error": str(e)}

    def cortex_compute_damages(self, lane=None):
        eng = self._get_cortex()
        if not eng:
            return {"error": "Cortex engine not available"}
        try:
            return eng.compute_damages(lane=lane)
        except Exception as e:
            return {"error": str(e)}

    def cortex_filing_qa(self, filing_id):
        eng = self._get_cortex()
        if not eng:
            return {"error": "Cortex engine not available"}
        try:
            return eng.filing_qa(filing_id)
        except Exception as e:
            return {"error": str(e)}

    def cortex_build_narrative(self, lane=None, limit=50):
        eng = self._get_cortex()
        if not eng:
            return {"error": "Cortex engine not available"}
        try:
            return eng.build_narrative(lane=lane, limit=limit)
        except Exception as e:
            return {"error": str(e)}

    def cortex_build_xexam(self, witness, topic=None):
        eng = self._get_cortex()
        if not eng:
            return {"error": "Cortex engine not available"}
        try:
            return eng.build_xexam(witness, topic=topic)
        except Exception as e:
            return {"error": str(e)}

    def cortex_trace_authority(self, citation):
        eng = self._get_cortex()
        if not eng:
            return {"error": "Cortex engine not available"}
        try:
            return eng.trace_authority(citation)
        except Exception as e:
            return {"error": str(e)}

    def cortex_lane_readiness(self, lane=None):
        eng = self._get_cortex()
        if not eng:
            return {"error": "Cortex engine not available"}
        try:
            return eng.lane_readiness(lane=lane)
        except Exception as e:
            return {"error": str(e)}

    def cortex_adversary_profile(self, adversary):
        eng = self._get_cortex()
        if not eng:
            return {"error": "Cortex engine not available"}
        try:
            return eng.adversary_profile(adversary)
        except Exception as e:
            return {"error": str(e)}

    def cortex_check_deadlines(self, days_ahead=30):
        eng = self._get_cortex()
        if not eng:
            return {"error": "Cortex engine not available"}
        try:
            return eng.check_deadlines(days_ahead=days_ahead)
        except Exception as e:
            return {"error": str(e)}

    def cortex_detect_gaps(self, lane=None):
        eng = self._get_cortex()
        if not eng:
            return {"error": "Cortex engine not available"}
        try:
            return eng.detect_gaps(lane=lane)
        except Exception as e:
            return {"error": str(e)}

    def cortex_generate_rebuttal(self, claim):
        eng = self._get_cortex()
        if not eng:
            return {"error": "Cortex engine not available"}
        try:
            return eng.generate_rebuttal(claim)
        except Exception as e:
            return {"error": str(e)}

    def cortex_nexus_fuse(self, topic, lane=None):
        eng = self._get_cortex()
        if not eng:
            return {"error": "Cortex engine not available"}
        try:
            return eng.nexus_fuse(topic, lane=lane)
        except Exception as e:
            return {"error": str(e)}

    def cortex_strategic_report(self, lane=None):
        eng = self._get_cortex()
        if not eng:
            return {"error": "Cortex engine not available"}
        try:
            return eng.strategic_report(lane=lane)
        except Exception as e:
            return {"error": str(e)}

    def cortex_search_all(self, query, limit=20):
        eng = self._get_cortex()
        if not eng:
            return {"error": "Cortex engine not available"}
        try:
            return eng.search_all(query, limit=limit)
        except Exception as e:
            return {"error": str(e)}

    # ===================================================================
    # AUTOMATON API — Legal Reasoning AGI Engine
    # ===================================================================

    def automaton_start(self, template_id=None):
        eng = self._get_automaton()
        if not eng:
            return {"error": "Automaton engine not available"}
        try:
            return eng.start(template_id=template_id)
        except Exception as e:
            return {"error": str(e)}

    def automaton_stop(self):
        eng = self._get_automaton()
        if not eng:
            return {"error": "Automaton engine not available"}
        try:
            return eng.stop()
        except Exception as e:
            return {"error": str(e)}

    def automaton_status(self):
        eng = self._get_automaton()
        if not eng:
            return {"error": "Automaton engine not available"}
        try:
            return eng.status()
        except Exception as e:
            return {"error": str(e)}

    def automaton_results(self):
        eng = self._get_automaton()
        if not eng:
            return {"error": "Automaton engine not available"}
        try:
            return eng.get_results()
        except Exception as e:
            return {"error": str(e)}

    def automaton_templates(self):
        eng = self._get_automaton()
        if not eng:
            return {"error": "Automaton engine not available"}
        try:
            return eng.get_templates()
        except Exception as e:
            return {"error": str(e)}

    def automaton_centrality(self):
        eng = self._get_automaton()
        if not eng:
            return {"error": "Automaton engine not available"}
        try:
            return eng.get_centrality()
        except Exception as e:
            return {"error": str(e)}

    # ===================================================================
    # BACKEND BRIDGE API — Filing, Drafting, Forms, PDF Tools
    # ===================================================================

    def backend_capabilities(self):
        eng = self._get_backend()
        if not eng:
            return {"error": "Backend bridge not available"}
        try:
            return eng.capabilities()
        except Exception as e:
            return {"error": str(e)}

    def backend_fraud_detect(self, csv_path=None, rows=None):
        eng = self._get_backend()
        if not eng:
            return {"error": "Backend bridge not available"}
        try:
            return eng.fraud_detect(csv_path=csv_path, rows=rows)
        except Exception as e:
            return {"error": str(e)}

    def backend_draft_complaint(self, kind="general", facts=None):
        eng = self._get_backend()
        if not eng:
            return {"error": "Backend bridge not available"}
        try:
            return eng.draft_complaint(kind=kind, facts=facts or {})
        except Exception as e:
            return {"error": str(e)}

    def backend_draft_injunction(self, purpose="", facts=None):
        eng = self._get_backend()
        if not eng:
            return {"error": "Backend bridge not available"}
        try:
            return eng.draft_injunction(purpose=purpose, facts=facts or {})
        except Exception as e:
            return {"error": str(e)}

    def backend_draft_show_cause(self, topic="", facts=None):
        eng = self._get_backend()
        if not eng:
            return {"error": "Backend bridge not available"}
        try:
            return eng.draft_show_cause(topic=topic, facts=facts or {})
        except Exception as e:
            return {"error": str(e)}

    def backend_draft_discovery(self, facts=None):
        eng = self._get_backend()
        if not eng:
            return {"error": "Backend bridge not available"}
        try:
            return eng.draft_discovery(facts=facts or {})
        except Exception as e:
            return {"error": str(e)}

    def backend_draft_disqualification(self, facts=None):
        eng = self._get_backend()
        if not eng:
            return {"error": "Backend bridge not available"}
        try:
            return eng.draft_disqualification(facts=facts or {})
        except Exception as e:
            return {"error": str(e)}

    def backend_draft_jtc(self, facts=None):
        eng = self._get_backend()
        if not eng:
            return {"error": "Backend bridge not available"}
        try:
            return eng.draft_jtc(facts=facts or {})
        except Exception as e:
            return {"error": str(e)}

    def backend_draft_appeal(self, facts=None):
        eng = self._get_backend()
        if not eng:
            return {"error": "Backend bridge not available"}
        try:
            return eng.draft_appeal(facts=facts or {})
        except Exception as e:
            return {"error": str(e)}

    def backend_draft_superintending(self, facts=None):
        eng = self._get_backend()
        if not eng:
            return {"error": "Backend bridge not available"}
        try:
            return eng.draft_superintending(facts=facts or {})
        except Exception as e:
            return {"error": str(e)}

    def backend_draft_relief(self, facts=None):
        eng = self._get_backend()
        if not eng:
            return {"error": "Backend bridge not available"}
        try:
            return eng.draft_relief(facts=facts or {})
        except Exception as e:
            return {"error": str(e)}

    def backend_bates_manifest(self, file_paths=None, prefix="PIGORS", start=1):
        eng = self._get_backend()
        if not eng:
            return {"error": "Backend bridge not available"}
        try:
            return eng.bates_manifest(file_paths=file_paths or [], prefix=prefix, start=start)
        except Exception as e:
            return {"error": str(e)}

    def backend_rules_tactical(self, rule_changes=None):
        eng = self._get_backend()
        if not eng:
            return {"error": "Backend bridge not available"}
        try:
            return eng.rules_tactical(rule_changes=rule_changes or [])
        except Exception as e:
            return {"error": str(e)}

    def backend_pdf_stamp(self, src_pdf="", out_pdf="", bates_text="", exhibit_text=""):
        eng = self._get_backend()
        if not eng:
            return {"error": "Backend bridge not available"}
        try:
            return eng.pdf_stamp(src_pdf=src_pdf, out_pdf=out_pdf,
                                 bates_text=bates_text, exhibit_text=exhibit_text)
        except Exception as e:
            return {"error": str(e)}

    def backend_pdf_normalize(self, src_pdf="", out_pdf=""):
        eng = self._get_backend()
        if not eng:
            return {"error": "Backend bridge not available"}
        try:
            return eng.pdf_normalize(src_pdf=src_pdf, out_pdf=out_pdf)
        except Exception as e:
            return {"error": str(e)}

    def backend_forms_resolve(self, action=""):
        eng = self._get_backend()
        if not eng:
            return {"error": "Backend bridge not available"}
        try:
            return eng.forms_resolve(action=action)
        except Exception as e:
            return {"error": str(e)}

    def backend_forms_catalog(self):
        eng = self._get_backend()
        if not eng:
            return {"error": "Backend bridge not available"}
        try:
            return eng.forms_catalog()
        except Exception as e:
            return {"error": str(e)}

    # ===================================================================
    # BLEEDING EDGE — DuckDB Analytics / LanceDB Semantic / Hybrid Search
    #                  GRAPHFORGE Visualization / NIGHTWATCH File Intel
    # ===================================================================

    def bleeding_edge_capabilities(self):
        eng = self._get_bleeding_edge()
        if not eng:
            return {"error": "Bleeding edge bridge not available"}
        try:
            return eng.capabilities()
        except Exception as e:
            return {"error": str(e)}

    # -- DuckDB Analytics --

    def analytics_evidence_heatmap(self):
        eng = self._get_bleeding_edge()
        if not eng:
            return {"error": "Bleeding edge bridge not available"}
        try:
            return eng.evidence_heatmap()
        except Exception as e:
            return {"error": str(e)}

    def analytics_authority_coverage(self):
        eng = self._get_bleeding_edge()
        if not eng:
            return {"error": "Bleeding edge bridge not available"}
        try:
            return eng.authority_coverage()
        except Exception as e:
            return {"error": str(e)}

    def analytics_impeachment_arsenal(self, target=None):
        eng = self._get_bleeding_edge()
        if not eng:
            return {"error": "Bleeding edge bridge not available"}
        try:
            return eng.impeachment_arsenal(target)
        except Exception as e:
            return {"error": str(e)}

    def analytics_filing_readiness(self):
        eng = self._get_bleeding_edge()
        if not eng:
            return {"error": "Bleeding edge bridge not available"}
        try:
            return eng.filing_readiness()
        except Exception as e:
            return {"error": str(e)}

    def analytics_judicial_dashboard(self):
        eng = self._get_bleeding_edge()
        if not eng:
            return {"error": "Bleeding edge bridge not available"}
        try:
            return eng.judicial_dashboard()
        except Exception as e:
            return {"error": str(e)}

    def analytics_contradiction_network(self, min_severity="high"):
        eng = self._get_bleeding_edge()
        if not eng:
            return {"error": "Bleeding edge bridge not available"}
        try:
            return eng.contradiction_network(min_severity)
        except Exception as e:
            return {"error": str(e)}

    def analytics_separation_counter(self):
        eng = self._get_bleeding_edge()
        if not eng:
            return {"error": "Bleeding edge bridge not available"}
        try:
            return eng.separation_counter()
        except Exception as e:
            return {"error": str(e)}

    def analytics_full_dashboard(self):
        eng = self._get_bleeding_edge()
        if not eng:
            return {"error": "Bleeding edge bridge not available"}
        try:
            return eng.full_dashboard()
        except Exception as e:
            return {"error": str(e)}

    # -- LanceDB Semantic Search --

    def semantic_search(self, query, table_name=None, top_k=10):
        eng = self._get_bleeding_edge()
        if not eng:
            return {"error": "Bleeding edge bridge not available"}
        try:
            return eng.semantic_search(query, table_name, top_k)
        except Exception as e:
            return {"error": str(e)}

    def semantic_encode(self, texts):
        eng = self._get_bleeding_edge()
        if not eng:
            return {"error": "Bleeding edge bridge not available"}
        try:
            return eng.semantic_encode(texts)
        except Exception as e:
            return {"error": str(e)}

    def semantic_lance_stats(self):
        eng = self._get_bleeding_edge()
        if not eng:
            return {"error": "Bleeding edge bridge not available"}
        try:
            return eng.lance_stats()
        except Exception as e:
            return {"error": str(e)}

    # -- Hybrid 3-Stage Search (FTS5 → LanceDB → CrossEncoder) --

    def hybrid_search(self, query, top_k=10, use_reranker=True):
        eng = self._get_bleeding_edge()
        if not eng:
            return {"error": "Bleeding edge bridge not available"}
        try:
            return eng.hybrid_search(query, top_k, use_reranker)
        except Exception as e:
            return {"error": str(e)}

    # -- GRAPHFORGE Visualization --

    def graphforge_get_palette(self, name="neon"):
        eng = self._get_bleeding_edge()
        if not eng:
            return {"error": "Bleeding edge bridge not available"}
        try:
            return eng.get_palette(name)
        except Exception as e:
            return {"error": str(e)}

    def graphforge_list_palettes(self):
        eng = self._get_bleeding_edge()
        if not eng:
            return {"error": "Bleeding edge bridge not available"}
        try:
            return eng.list_palettes()
        except Exception as e:
            return {"error": str(e)}

    def graphforge_scan_database(self, db_path=None):
        eng = self._get_bleeding_edge()
        if not eng:
            return {"error": "Bleeding edge bridge not available"}
        try:
            return eng.scan_database(db_path)
        except Exception as e:
            return {"error": str(e)}

    def graphforge_merge_databases(self, db_paths=None):
        eng = self._get_bleeding_edge()
        if not eng:
            return {"error": "Bleeding edge bridge not available"}
        try:
            return eng.merge_databases(db_paths or [])
        except Exception as e:
            return {"error": str(e)}

    def graphforge_export_json(self):
        eng = self._get_bleeding_edge()
        if not eng:
            return {"error": "Bleeding edge bridge not available"}
        try:
            return eng.export_graph_json()
        except Exception as e:
            return {"error": str(e)}

    def graphforge_export_csv(self):
        eng = self._get_bleeding_edge()
        if not eng:
            return {"error": "Bleeding edge bridge not available"}
        try:
            return eng.export_graph_csv()
        except Exception as e:
            return {"error": str(e)}

    def graphforge_build_csv_graph(self, csv_path, id_col="id", label_col="label",
                                   group_col=None, link_cols=None, palette_name="neon"):
        eng = self._get_bleeding_edge()
        if not eng:
            return {"error": "Bleeding edge bridge not available"}
        try:
            return eng.build_csv_graph(csv_path, id_col, label_col,
                                       group_col, link_cols, palette_name)
        except Exception as e:
            return {"error": str(e)}

    # -- NIGHTWATCH File Intelligence --

    def nightwatch_hash_file(self, filepath):
        eng = self._get_bleeding_edge()
        if not eng:
            return {"error": "Bleeding edge bridge not available"}
        try:
            return eng.hash_file(filepath)
        except Exception as e:
            return {"error": str(e)}

    def nightwatch_scan_directory(self, directory, extensions=None):
        eng = self._get_bleeding_edge()
        if not eng:
            return {"error": "Bleeding edge bridge not available"}
        try:
            return eng.scan_directory_changes(directory, extensions)
        except Exception as e:
            return {"error": str(e)}

    # ===================================================================
    # INTELLIGENCE ENGINE API (11 engines, 75+ methods via IntelligenceBridge)
    # ===================================================================

    def _intel(self, method, *args, **kwargs):
        """Generic intelligence bridge dispatcher."""
        eng = self._get_intelligence()
        if not eng:
            return {"error": "Intelligence bridge not available"}
        fn = getattr(eng, method, None)
        if fn is None:
            return {"error": f"Unknown intelligence method: {method}"}
        try:
            return fn(*args, **kwargs)
        except Exception as e:
            return {"error": str(e)}

    def intelligence_status(self):
        eng = self._get_intelligence()
        return eng.status() if eng else {"error": "Intelligence bridge not available"}

    # -- NEXUS Intelligence --
    def nexus_fuse(self, topic, lanes=None, limit=50):
        return self._intel("nexus_fuse", topic, lanes, limit)

    def nexus_argue(self, claim, lane=None, limit=10):
        return self._intel("nexus_argue", claim, lane, limit)

    def nexus_readiness(self, lane=None):
        return self._intel("nexus_readiness", lane)

    def nexus_damages(self, lane=None):
        return self._intel("nexus_damages", lane)

    def nexus_red_team(self, claim):
        return self._intel("nexus_red_team", claim)

    def nexus_case_map(self, case_type):
        return self._intel("nexus_case_map", case_type)

    def nexus_priorities(self):
        return self._intel("nexus_priorities")

    # -- NEMESIS Adversary Intelligence --
    def nemesis_profile(self, target):
        return self._intel("nemesis_profile", target)

    def nemesis_predict(self, target):
        return self._intel("nemesis_predict", target)

    def nemesis_vulnerabilities(self, target):
        return self._intel("nemesis_vulnerabilities", target)

    # -- CERBERUS Filing Validation --
    def cerberus_report(self, lane):
        return self._intel("cerberus_report", lane)

    def cerberus_gaps(self, lane):
        return self._intel("cerberus_gaps", lane)

    def cerberus_weapons(self, lane=None):
        return self._intel("cerberus_weapons", lane)

    # -- CHIMERA Multi-Source Blending --
    def chimera_detect(self, target):
        return self._intel("chimera_detect", target)

    def chimera_impeachment(self, target):
        return self._intel("chimera_impeachment", target)

    def chimera_patterns(self, target=None):
        return self._intel("chimera_patterns", target)

    # -- TEMPORAL Knowledge Graph --
    def temporal_poisonous_tree(self, event_id=None):
        return self._intel("temporal_poisonous_tree", event_id)

    def temporal_conspiracy(self):
        return self._intel("temporal_conspiracy")

    def temporal_anomalies(self):
        return self._intel("temporal_anomalies")

    # -- CAUSAL Chain Engine --
    def causal_poisonous_tree(self, root_event=None):
        return self._intel("causal_poisonous_tree", root_event)

    def causal_retaliation(self):
        return self._intel("causal_retaliation")

    def causal_for_filing(self, lane=None):
        return self._intel("causal_for_filing", lane)

    # -- DAMAGES Engine --
    def damages_total(self):
        return self._intel("damages_total")

    def damages_by_lane(self, lane=None):
        return self._intel("damages_by_lane", lane)

    def damages_summary(self):
        return self._intel("damages_summary")

    # -- IRAC Argument Engine --
    def irac_argument(self, claim, lane=None):
        return self._intel("irac_argument", claim, lane)

    def irac_strong(self, lane=None, limit=10):
        return self._intel("irac_strong", lane, limit)

    def irac_gaps(self, lane=None):
        return self._intel("irac_gaps", lane)

    # -- HYPERGRAPH Evidence Engine --
    def hypergraph_conspiracy(self):
        return self._intel("hypergraph_conspiracy")

    def hypergraph_strongest(self, lane=None, limit=10):
        return self._intel("hypergraph_strongest", lane, limit)

    def hypergraph_filing_ammo(self, lane):
        return self._intel("hypergraph_filing_ammo", lane)

    # -- NARRATIVE Builder --
    def narrative_claim(self, claim):
        return self._intel("narrative_claim", claim)

    def narrative_defendant(self, defendant):
        return self._intel("narrative_defendant", defendant)

    def narrative_exhibits(self, lane=None):
        return self._intel("narrative_exhibits", lane)

    # -- FORGE Filing Engine --
    def forge_assemble(self, lane, filing_type=None):
        return self._intel("forge_assemble", lane, filing_type)

    def forge_validate(self, lane):
        return self._intel("forge_validate", lane)

    def forge_list_packets(self):
        return self._intel("forge_list_packets")

    # ===================================================================
    # OPERATIONS API (Cycle 6 — OperationsBridge: 10 engines, 26 methods)
    # ===================================================================

    def _ops(self, method, *args, **kwargs):
        eng = self._get_operations()
        if not eng:
            return {"error": "Operations bridge not available"}
        fn = getattr(eng, method, None)
        if fn is None:
            return {"error": f"Unknown operations method: {method}"}
        try:
            return fn(*args, **kwargs)
        except Exception as e:
            return {"error": str(e)}

    def operations_status(self):
        return self._ops("status")

    # -- Orchestrator --
    def orchestrator_fuse(self, query):
        return self._ops("orchestrator_fuse", query)

    def orchestrator_entity_profile(self, name):
        return self._ops("orchestrator_entity_profile", name)

    def orchestrator_filing_intel(self, lane):
        return self._ops("orchestrator_filing_intel", lane)

    def orchestrator_health(self):
        return self._ops("orchestrator_health")

    # -- QA Engine --
    def qa_audit_all(self):
        return self._ops("qa_audit_all")

    def qa_summary(self, report=None):
        return self._ops("qa_summary", report)

    # -- Filing Engine --
    def filing_scan_triggers(self):
        return self._ops("filing_scan_triggers")

    def filing_run(self, filing_id, case_num=None):
        return self._ops("filing_run", filing_id, case_num)

    def filing_engine_status(self):
        return self._ops("filing_engine_status")

    # -- Filing Assembler --
    def assembler_build_stack(self, filing_id=None):
        return self._ops("assembler_build_stack", filing_id)

    # -- Intake Analyzer --
    def analyze_document(self, text, file_name=None):
        return self._ops("analyze_document", text, file_name)

    # -- Rebuttal Engine --
    def rebuttal_all(self):
        return self._ops("rebuttal_all")

    def rebuttal_by_factor(self, factor):
        return self._ops("rebuttal_by_factor", factor)

    def rebuttal_appellate_section(self):
        return self._ops("rebuttal_appellate_section")

    # -- Perception (Legal-BERT) --
    def perception_classify(self, text):
        return self._ops("perception_classify", text)

    def perception_ner(self, text):
        return self._ops("perception_ner", text)

    def perception_embed(self, text):
        return self._ops("perception_embed", text)

    # -- Apotheosis (System Health) --
    def apotheosis_health(self):
        return self._ops("apotheosis_health")

    def apotheosis_gaps(self):
        return self._ops("apotheosis_gaps")

    def apotheosis_recommendations(self):
        return self._ops("apotheosis_recommendations")

    def apotheosis_trajectory(self):
        return self._ops("apotheosis_trajectory")

    # -- Lexicon --
    def lexicon_filing_roadmap(self):
        return self._ops("lexicon_filing_roadmap")

    def lexicon_compute_deadlines(self):
        return self._ops("lexicon_compute_deadlines")

    def lexicon_evidence_check(self):
        return self._ops("lexicon_evidence_check")

    def lexicon_canon_analysis(self):
        return self._ops("lexicon_canon_analysis")

    def lexicon_ask(self, question):
        return self._ops("lexicon_ask", question)

    # -----------------------------------------------------------------------
    # WAR ROOM API methods — called by index.html War Room JS
    # -----------------------------------------------------------------------

    def adversary_profile(self, name):
        """Deep adversary intelligence — direct DB queries (no Cortex dependency)."""
        conn = _connect_lit()
        if not conn:
            return {"name": name, "error": "DB unavailable"}
        like = f"%{name[:100]}%"
        try:
            result = {"name": name}
            # Evidence count
            ev = 0
            if _table_exists(conn, "evidence_quotes"):
                ev = conn.execute(
                    "SELECT COUNT(*) FROM evidence_quotes WHERE quote_text LIKE ? AND is_duplicate=0", (like,)
                ).fetchone()[0]
            result["evidence_count"] = ev
            # Impeachment count
            imp = 0
            if _table_exists(conn, "impeachment_matrix"):
                imp = conn.execute(
                    "SELECT COUNT(*) FROM impeachment_matrix WHERE evidence_summary LIKE ? OR quote_text LIKE ?",
                    (like, like),
                ).fetchone()[0]
            result["impeachment_count"] = imp
            # Contradiction count
            con = 0
            if _table_exists(conn, "contradiction_map"):
                con = conn.execute(
                    "SELECT COUNT(*) FROM contradiction_map WHERE source_a LIKE ? OR source_b LIKE ? OR contradiction_text LIKE ?",
                    (like, like, like),
                ).fetchone()[0]
            result["contradiction_count"] = con
            # Timeline count
            tl = 0
            if _table_exists(conn, "timeline_events"):
                tl = conn.execute(
                    "SELECT COUNT(*) FROM timeline_events WHERE event_description LIKE ? OR actor LIKE ?",
                    (like, like),
                ).fetchone()[0]
            result["timeline_count"] = tl
            # Judicial violations
            viol = 0
            if _table_exists(conn, "judicial_violations"):
                viol = conn.execute(
                    "SELECT COUNT(*) FROM judicial_violations WHERE description LIKE ?", (like,)
                ).fetchone()[0]
            result["violation_count"] = viol
            # Cartel connections
            cartel = []
            if _table_exists(conn, "berry_mcneill_intelligence"):
                rows = conn.execute(
                    "SELECT person_a, person_b, connection_type, notes FROM berry_mcneill_intelligence "
                    "WHERE person_a LIKE ? OR person_b LIKE ? LIMIT 10",
                    (like, like),
                ).fetchall()
                cartel = [{"person_a": r[0], "person_b": r[1], "type": r[2], "notes": r[3]} for r in rows]
            result["cartel"] = cartel
            # Lane breakdown
            lane_breakdown = {}
            if _table_exists(conn, "evidence_quotes"):
                rows = conn.execute(
                    "SELECT lane, COUNT(*) FROM evidence_quotes WHERE quote_text LIKE ? AND is_duplicate=0 AND lane IS NOT NULL GROUP BY lane",
                    (like,),
                ).fetchall()
                lane_breakdown = {r[0]: r[1] for r in rows}
            result["lane_breakdown"] = lane_breakdown
            # Top quotes
            top_quotes = []
            if _table_exists(conn, "evidence_quotes"):
                rows = conn.execute(
                    "SELECT quote_text, source_file, category, lane FROM evidence_quotes "
                    "WHERE quote_text LIKE ? AND is_duplicate=0 ORDER BY CAST(relevance_score AS REAL) DESC LIMIT 5",
                    (like,),
                ).fetchall()
                top_quotes = [{"quote": r[0][:200], "source": r[1], "category": r[2], "lane": r[3]} for r in rows]
            result["top_quotes"] = top_quotes
            # Threat score (weighted composite)
            result["threat_score"] = min(100, int(ev * 0.02 + imp * 0.5 + con * 0.8 + viol * 0.3 + tl * 0.01))
            result["separation_days"] = _sep_days()
            return result
        except Exception as e:
            return {"name": name, "error": str(e)}
        finally:
            conn.close()

    def search_impeachment_matrix(self, target="", category="", min_severity=1, limit=20):
        """Search impeachment_matrix for cross-examination ammo."""
        conn = _connect_lit()
        if not conn or not _table_exists(conn, "impeachment_matrix"):
            conn and conn.close()
            return []
        try:
            clauses, params = [], []
            if target:
                clauses.append("(evidence_summary LIKE ? OR category LIKE ?)")
                t = f"%{target[:100]}%"
                params.extend([t, t])
            if category:
                clauses.append("category LIKE ?")
                params.append(f"%{category[:60]}%")
            if min_severity and int(min_severity) > 1:
                clauses.append("CAST(impeachment_value AS INTEGER) >= ?")
                params.append(int(min_severity))
            where = " AND ".join(clauses) if clauses else "1=1"
            sql = (
                f"SELECT category, evidence_summary, source_file, quote_text, "
                f"impeachment_value, cross_exam_question, filing_relevance, event_date "
                f"FROM impeachment_matrix WHERE {where} "
                f"ORDER BY CAST(impeachment_value AS INTEGER) DESC LIMIT ?"
            )
            params.append(int(limit))
            rows = conn.execute(sql, tuple(params)).fetchall()
            return _rows_to_dicts(rows)
        except Exception as e:
            return [{"error": str(e)}]
        finally:
            conn.close()

    def get_war_room_dossier_detail(self, name):
        """Extended dossier for an adversary — evidence + violations + cartel + impeachment + contradictions + harms."""
        conn = _connect_lit()
        if not conn:
            return {"name": name, "error": "DB unavailable"}
        like = f"%{name[:100]}%"
        result = {"name": name, "separation_days": _sep_days()}
        try:
            if _table_exists(conn, "evidence_quotes"):
                result["evidence_count"] = conn.execute(
                    "SELECT COUNT(*) FROM evidence_quotes WHERE is_duplicate=0 AND quote_text LIKE ?",
                    (like,),
                ).fetchone()[0]
                result["top_evidence"] = _rows_to_dicts(conn.execute(
                    "SELECT quote_text, source_file, category, lane FROM evidence_quotes "
                    "WHERE is_duplicate=0 AND quote_text LIKE ? "
                    "ORDER BY CAST(relevance_score AS REAL) DESC LIMIT 10", (like,)
                ).fetchall())
                result["lane_breakdown"] = _rows_to_dicts(conn.execute(
                    "SELECT lane, COUNT(*) as cnt FROM evidence_quotes "
                    "WHERE is_duplicate=0 AND quote_text LIKE ? GROUP BY lane ORDER BY cnt DESC",
                    (like,),
                ).fetchall())
            if _table_exists(conn, "judicial_violations"):
                result["violation_count"] = conn.execute(
                    "SELECT COUNT(*) FROM judicial_violations WHERE description LIKE ?",
                    (like,),
                ).fetchone()[0]
                result["violation_types"] = _rows_to_dicts(conn.execute(
                    "SELECT violation_type, COUNT(*) as cnt FROM judicial_violations "
                    "WHERE description LIKE ? GROUP BY violation_type ORDER BY cnt DESC LIMIT 10",
                    (like,),
                ).fetchall())
            if _table_exists(conn, "berry_mcneill_intelligence"):
                result["cartel"] = _rows_to_dicts(conn.execute(
                    "SELECT person_a, person_b, connection_type, relationship, notes "
                    "FROM berry_mcneill_intelligence "
                    "WHERE person_a LIKE ? OR person_b LIKE ? LIMIT 10", (like, like)
                ).fetchall())
            if _table_exists(conn, "police_reports"):
                result["police_reports"] = _rows_to_dicts(conn.execute(
                    "SELECT incident_numbers, dates, allegations, exculpatory, key_quotes "
                    "FROM police_reports "
                    "WHERE allegations LIKE ? OR exculpatory LIKE ? OR key_quotes LIKE ? LIMIT 10",
                    (like, like, like),
                ).fetchall())
            if _table_exists(conn, "impeachment_matrix"):
                result["impeachment_count"] = conn.execute(
                    "SELECT COUNT(*) FROM impeachment_matrix "
                    "WHERE evidence_summary LIKE ? OR quote_text LIKE ?",
                    (like, like),
                ).fetchone()[0]
                result["top_impeachment"] = _rows_to_dicts(conn.execute(
                    "SELECT category, evidence_summary, impeachment_value, cross_exam_question "
                    "FROM impeachment_matrix "
                    "WHERE evidence_summary LIKE ? OR quote_text LIKE ? "
                    "ORDER BY CAST(impeachment_value AS INTEGER) DESC LIMIT 5",
                    (like, like),
                ).fetchall())
            if _table_exists(conn, "contradiction_map"):
                result["contradiction_count"] = conn.execute(
                    "SELECT COUNT(*) FROM contradiction_map "
                    "WHERE source_a LIKE ? OR source_b LIKE ? OR contradiction_text LIKE ?",
                    (like, like, like),
                ).fetchone()[0]
                result["top_contradictions"] = _rows_to_dicts(conn.execute(
                    "SELECT source_a, source_b, contradiction_text, severity "
                    "FROM contradiction_map "
                    "WHERE source_a LIKE ? OR source_b LIKE ? OR contradiction_text LIKE ? "
                    "ORDER BY CASE severity WHEN 'critical' THEN 1 WHEN 'high' THEN 2 ELSE 3 END "
                    "LIMIT 5",
                    (like, like, like),
                ).fetchall())
            if _table_exists(conn, "timeline_events"):
                result["timeline_count"] = conn.execute(
                    "SELECT COUNT(*) FROM timeline_events WHERE event_description LIKE ? OR actor LIKE ?",
                    (like, like),
                ).fetchone()[0]
            if _table_exists(conn, "extracted_harms"):
                result["harms"] = _rows_to_dicts(conn.execute(
                    "SELECT category, subcategory, description, severity, lane "
                    "FROM extracted_harms WHERE adversary LIKE ? OR description LIKE ? "
                    "ORDER BY CASE severity WHEN 'critical' THEN 1 WHEN 'high' THEN 2 ELSE 3 END "
                    "LIMIT 10",
                    (like, like),
                ).fetchall())
            # Compute threat score
            ev = result.get("evidence_count", 0)
            imp = result.get("impeachment_count", 0)
            con = result.get("contradiction_count", 0)
            viol = result.get("violation_count", 0)
            result["threat_score"] = min(100, int((ev * 0.02) + (imp * 0.5) + (con * 1.5) + (viol * 0.1)))
        except Exception as e:
            result["error"] = str(e)
        finally:
            conn.close()
        return result

    def timeline_search(self, query="", date_from=None, date_to=None, actor=None, limit=50):
        """Search timeline_events with FTS5 safety."""
        conn = _connect_lit()
        if not conn:
            return []
        try:
            safe_q = _sanitize_fts(str(query)) if query else ""
            if safe_q and _table_exists(conn, "timeline_fts"):
                try:
                    rows = conn.execute(
                        "SELECT t.event_date, t.event_description, t.lane, t.actor "
                        "FROM timeline_fts f JOIN timeline_events t ON f.rowid = t.rowid "
                        "WHERE timeline_fts MATCH ? LIMIT ?",
                        (safe_q, int(limit)),
                    ).fetchall()
                    if rows:
                        return _rows_to_dicts(rows)
                except Exception:
                    pass
            clauses, params = [], []
            if safe_q:
                clauses.append("event_description LIKE ?")
                params.append(f"%{safe_q}%")
            if date_from:
                clauses.append("event_date >= ?")
                params.append(str(date_from))
            if date_to:
                clauses.append("event_date <= ?")
                params.append(str(date_to))
            if actor:
                clauses.append("actor LIKE ?")
                params.append(f"%{actor[:80]}%")
            where = " AND ".join(clauses) if clauses else "1=1"
            sql = (
                f"SELECT event_date, event_description, lane, actor "
                f"FROM timeline_events WHERE {where} "
                f"ORDER BY event_date DESC LIMIT ?"
            )
            params.append(int(limit))
            return _rows_to_dicts(conn.execute(sql, tuple(params)).fetchall())
        except Exception as e:
            return [{"error": str(e)}]
        finally:
            conn.close()

    def search_evidence_quotes(self, query):
        """FTS5 search on evidence_quotes with LIKE fallback."""
        conn = _connect_lit()
        if not conn:
            return []
        try:
            safe_q = _sanitize_fts(str(query)) if query else ""
            if not safe_q:
                return []
            if _table_exists(conn, "evidence_fts"):
                try:
                    rows = conn.execute(
                        "SELECT e.quote_text, e.source_file, e.category, e.lane, e.page_number "
                        "FROM evidence_fts f JOIN evidence_quotes e ON f.rowid = e.rowid "
                        "WHERE evidence_fts MATCH ? LIMIT 30",
                        (safe_q,),
                    ).fetchall()
                    if rows:
                        return _rows_to_dicts(rows)
                except Exception:
                    pass
            like_pat = f"%{safe_q}%"
            return _rows_to_dicts(conn.execute(
                "SELECT quote_text, source_file, category, lane, page_number "
                "FROM evidence_quotes WHERE quote_text LIKE ? AND is_duplicate=0 LIMIT 30",
                (like_pat,),
            ).fetchall())
        except Exception as e:
            return [{"error": str(e)}]
        finally:
            conn.close()

    def search_authority_chains(self, citation):
        """Search authority_chains_v2 by citation."""
        conn = _connect_lit()
        if not conn or not _table_exists(conn, "authority_chains_v2"):
            conn and conn.close()
            return []
        try:
            like = f"%{str(citation)[:100]}%"
            return _rows_to_dicts(conn.execute(
                "SELECT primary_citation, supporting_citation, relationship, "
                "source_document, lane "
                "FROM authority_chains_v2 "
                "WHERE primary_citation LIKE ? OR supporting_citation LIKE ? "
                "ORDER BY primary_citation LIMIT 30",
                (like, like),
            ).fetchall())
        except Exception as e:
            return [{"error": str(e)}]
        finally:
            conn.close()

    def search_contradictions(self, entity):
        """Search contradiction_map for inconsistencies."""
        conn = _connect_lit()
        if not conn or not _table_exists(conn, "contradiction_map"):
            conn and conn.close()
            return []
        try:
            like = f"%{str(entity)[:100]}%"
            return _rows_to_dicts(conn.execute(
                "SELECT claim_id, source_a, source_b, contradiction_text, severity, lane "
                "FROM contradiction_map "
                "WHERE source_a LIKE ? OR source_b LIKE ? OR contradiction_text LIKE ? "
                "ORDER BY CASE severity WHEN 'critical' THEN 1 WHEN 'high' THEN 2 ELSE 3 END "
                "LIMIT 30",
                (like, like, like),
            ).fetchall())
        except Exception as e:
            return [{"error": str(e)}]
        finally:
            conn.close()

    def filing_readiness(self, lane=None):
        """Filing readiness dashboard — optionally filtered by lane."""
        conn = _connect_lit()
        if not conn:
            return {"error": "DB unavailable"}
        result = {"lane": lane, "separation_days": _sep_days()}
        try:
            if _table_exists(conn, "filing_packages"):
                where = "WHERE lane = ?" if lane else ""
                params = (str(lane),) if lane else ()
                result["filings"] = _rows_to_dicts(conn.execute(
                    f"SELECT filing_id, title, lane, case_number, status, doc_count "
                    f"FROM filing_packages {where} "
                    f"ORDER BY doc_count DESC LIMIT 20",
                    params,
                ).fetchall())
            if _table_exists(conn, "evidence_quotes"):
                where = "WHERE lane = ? AND is_duplicate=0" if lane else "WHERE is_duplicate=0"
                params = (str(lane),) if lane else ()
                result["evidence_count"] = conn.execute(
                    f"SELECT COUNT(*) FROM evidence_quotes {where}", params
                ).fetchone()[0]
                # Per-lane breakdown
                result["evidence_by_lane"] = _rows_to_dicts(conn.execute(
                    "SELECT lane, COUNT(*) as cnt FROM evidence_quotes "
                    "WHERE is_duplicate=0 AND lane IS NOT NULL GROUP BY lane ORDER BY cnt DESC"
                ).fetchall())
            if _table_exists(conn, "authority_chains_v2"):
                where = "WHERE lane = ?" if lane else ""
                params = (str(lane),) if lane else ()
                result["authority_count"] = conn.execute(
                    f"SELECT COUNT(*) FROM authority_chains_v2 {where}", params
                ).fetchone()[0]
            if _table_exists(conn, "impeachment_matrix"):
                result["impeachment_count"] = conn.execute(
                    "SELECT COUNT(*) FROM impeachment_matrix"
                ).fetchone()[0]
            if _table_exists(conn, "contradiction_map"):
                where = "WHERE lane = ?" if lane else ""
                params = (str(lane),) if lane else ()
                result["contradiction_count"] = conn.execute(
                    f"SELECT COUNT(*) FROM contradiction_map {where}", params
                ).fetchone()[0]
            if _table_exists(conn, "extracted_harms"):
                where = "WHERE lane = ?" if lane else ""
                params = (str(lane),) if lane else ()
                result["harms_count"] = conn.execute(
                    f"SELECT COUNT(*) FROM extracted_harms {where}", params
                ).fetchone()[0]
            if _table_exists(conn, "deadlines"):
                today_str = str(date.today())
                result["upcoming_deadlines"] = _rows_to_dicts(conn.execute(
                    "SELECT title, due_date, status, court FROM deadlines "
                    "WHERE due_date >= ? ORDER BY due_date LIMIT 10",
                    (today_str,),
                ).fetchall())
            # Readiness score: evidence + authorities + impeachment + contradictions
            ev = result.get("evidence_count", 0)
            auth = result.get("authority_count", 0)
            imp = result.get("impeachment_count", 0)
            con = result.get("contradiction_count", 0)
            score = min(100, int((min(ev, 500) / 5) + (min(auth, 200) / 2) + imp * 0.3 + con * 0.5))
            result["readiness_score"] = score
        except Exception as e:
            result["error"] = str(e)
        finally:
            conn.close()
        return result

    def get_war_room_situation(self):
        """Situation dashboard: deadlines + case status + separation counter + urgency."""
        conn = _connect_lit()
        today = date.today()
        today_str = str(today)
        result = {
            "separation_days": _sep_days(),
            "separation_date": str(SEPARATION_DATE),
            "today": today_str,
        }
        if not conn:
            return result
        try:
            if _table_exists(conn, "deadlines"):
                raw_deadlines = _rows_to_dicts(conn.execute(
                    "SELECT title, due_date, status, court, case_number "
                    "FROM deadlines ORDER BY due_date LIMIT 30"
                ).fetchall())
                # Add urgency flags
                for d in raw_deadlines:
                    dd = d.get("due_date", "")
                    if dd:
                        try:
                            diff = (date.fromisoformat(dd[:10]) - today).days
                            if diff < 0:
                                d["urgency"] = "overdue"
                            elif diff <= 3:
                                d["urgency"] = "critical"
                            elif diff <= 7:
                                d["urgency"] = "urgent"
                            else:
                                d["urgency"] = "ok"
                            d["days_until"] = diff
                        except (ValueError, TypeError):
                            d["urgency"] = "unknown"
                    d["text"] = d.get("title", "")
                    d["date"] = d.get("due_date", "")
                    d["description"] = f"{d.get('court', '')} — {d.get('title', '')}"
                result["deadlines"] = raw_deadlines
                # Aggregate urgency flags
                uf = {"overdue": 0, "critical": 0, "urgent": 0, "ok": 0}
                for d in raw_deadlines:
                    u = d.get("urgency", "unknown")
                    if u in uf:
                        uf[u] += 1
                result["urgency_flags"] = uf
            if _table_exists(conn, "filing_packages"):
                result["filings"] = _rows_to_dicts(conn.execute(
                    "SELECT filing_id, title, lane, status, doc_count, case_number "
                    "FROM filing_packages ORDER BY doc_count DESC LIMIT 20"
                ).fetchall())
            counts = {}
            for tbl in ("evidence_quotes", "authority_chains_v2", "timeline_events",
                        "impeachment_matrix", "contradiction_map", "judicial_violations",
                        "extracted_harms", "police_reports", "best_interest_factor_map",
                        "critical_facts"):
                if _table_exists(conn, tbl):
                    counts[tbl] = conn.execute(f"SELECT COUNT(*) FROM {tbl}").fetchone()[0]
            result["table_counts"] = counts
            # Per-lane evidence density (dict: {lane: count})
            if _table_exists(conn, "evidence_quotes"):
                ld_rows = conn.execute(
                    "SELECT lane, COUNT(*) as cnt "
                    "FROM evidence_quotes WHERE is_duplicate=0 AND lane IS NOT NULL "
                    "GROUP BY lane ORDER BY cnt DESC"
                ).fetchall()
                result["lane_density"] = {r[0]: r[1] for r in ld_rows if r[0]}
            # Overall readiness score
            total_ev = counts.get("evidence_quotes", 0)
            total_auth = counts.get("authority_chains_v2", 0)
            total_imp = counts.get("impeachment_matrix", 0)
            result["overall_readiness"] = min(100, int(
                (min(total_ev, 1000) / 10) + (min(total_auth, 500) / 5) + total_imp * 0.2
            ))
        except Exception as e:
            result["error"] = str(e)
        finally:
            conn.close()
        return result

    # -- War Room: Custody Factors ------------------------------------------
    def get_custody_factors(self):
        """MCL 722.23 best-interest factor analysis with evidence counts."""
        conn = _connect_lit()
        if not conn:
            return {"error": "DB unavailable"}
        result = {"factors": [], "total_evidence": 0}
        try:
            if _table_exists(conn, "best_interest_factor_map"):
                result["factors"] = _rows_to_dicts(conn.execute(
                    "SELECT factor, COUNT(*) as evidence_count, "
                    "GROUP_CONCAT(DISTINCT lane) as lanes "
                    "FROM best_interest_factor_map "
                    "GROUP BY factor ORDER BY evidence_count DESC"
                ).fetchall())
                result["total_evidence"] = conn.execute(
                    "SELECT COUNT(*) FROM best_interest_factor_map"
                ).fetchone()[0]
        except Exception as e:
            result["error"] = str(e)
        finally:
            conn.close()
        return result

    # -- War Room: Critical Facts -------------------------------------------
    def get_critical_facts(self, limit=50):
        """Verified immutable facts from critical_facts table."""
        conn = _connect_lit()
        if not conn:
            return {"error": "DB unavailable"}
        result = {"facts": [], "total": 0}
        try:
            if _table_exists(conn, "critical_facts"):
                cols = {r[1] for r in conn.execute("PRAGMA table_info(critical_facts)").fetchall()}
                select_cols = ", ".join(c for c in ["fact_text", "category", "source", "verified", "lane"] if c in cols) or "*"
                result["facts"] = _rows_to_dicts(conn.execute(
                    f"SELECT {select_cols} FROM critical_facts LIMIT ?", (int(limit),)
                ).fetchall())
                result["total"] = conn.execute("SELECT COUNT(*) FROM critical_facts").fetchone()[0]
        except Exception as e:
            result["error"] = str(e)
        finally:
            conn.close()
        return result

    # -- War Room: Extracted Harms ------------------------------------------
    def get_extracted_harms(self, lane=None, limit=30):
        """Damages evidence from extracted_harms table."""
        conn = _connect_lit()
        if not conn:
            return {"error": "DB unavailable"}
        result = {"harms": [], "total": 0, "by_category": []}
        try:
            if _table_exists(conn, "extracted_harms"):
                cols = {r[1] for r in conn.execute("PRAGMA table_info(extracted_harms)").fetchall()}
                where, params = "", ()
                if lane and "lane" in cols:
                    where, params = "WHERE lane = ?", (str(lane),)
                result["total"] = conn.execute(
                    f"SELECT COUNT(*) FROM extracted_harms {where}", params
                ).fetchone()[0]
                select_cols = ", ".join(c for c in ["category", "subcategory", "description", "severity", "lane", "adversary"] if c in cols) or "*"
                result["harms"] = _rows_to_dicts(conn.execute(
                    f"SELECT {select_cols} FROM extracted_harms {where} LIMIT ?",
                    params + (int(limit),),
                ).fetchall())
                if "category" in cols:
                    result["by_category"] = _rows_to_dicts(conn.execute(
                        f"SELECT category, COUNT(*) as cnt FROM extracted_harms {where} "
                        f"GROUP BY category ORDER BY cnt DESC",
                        params,
                    ).fetchall())
        except Exception as e:
            result["error"] = str(e)
        finally:
            conn.close()
        return result

    # -- War Room: Pre-built Graph Data ------------------------------------
    def get_graph_data_from_db(self, limit=500):
        """Load pre-built graph nodes+edges from graph_nodes/graph_edges tables."""
        conn = _connect_lit()
        if not conn:
            return {"error": "DB unavailable"}
        result = {"nodes": [], "edges": [], "node_count": 0, "edge_count": 0}
        try:
            if _table_exists(conn, "graph_nodes"):
                result["node_count"] = conn.execute("SELECT COUNT(*) FROM graph_nodes").fetchone()[0]
                cols = {r[1] for r in conn.execute("PRAGMA table_info(graph_nodes)").fetchall()}
                select_cols = ", ".join(c for c in ["id", "label", "layer", "group_name", "size", "color", "desc"] if c in cols) or "*"
                result["nodes"] = _rows_to_dicts(conn.execute(
                    f"SELECT {select_cols} FROM graph_nodes LIMIT ?", (int(limit),)
                ).fetchall())
            if _table_exists(conn, "graph_edges"):
                result["edge_count"] = conn.execute("SELECT COUNT(*) FROM graph_edges").fetchone()[0]
                cols = {r[1] for r in conn.execute("PRAGMA table_info(graph_edges)").fetchall()}
                select_cols = ", ".join(c for c in ["source", "target", "layer", "color", "weight", "type"] if c in cols) or "*"
                result["edges"] = _rows_to_dicts(conn.execute(
                    f"SELECT {select_cols} FROM graph_edges LIMIT ?", (int(limit * 2),)
                ).fetchall())
        except Exception as e:
            result["error"] = str(e)
        finally:
            conn.close()
        return result


# ---------------------------------------------------------------------------
# HTTP Server
# ---------------------------------------------------------------------------
class _QuietHandler(http.server.SimpleHTTPRequestHandler):
    """Serve V15 HTML assets with gzip compression and smart caching."""

    _gzip_cache: dict = {}  # path → (mtime, compressed_bytes)

    def __init__(self, *args, **kwargs):
        super().__init__(*args, directory=str(VIS_DIR), **kwargs)

    def log_message(self, fmt, *args):
        pass

    def end_headers(self):
        self.send_header("Access-Control-Allow-Origin", "*")
        super().end_headers()

    def guess_type(self, path):
        ext = os.path.splitext(path)[1].lower()
        overrides = {".json": "application/json", ".js": "application/javascript",
                     ".mjs": "application/javascript", ".woff2": "font/woff2"}
        return overrides.get(ext, super().guess_type(path))

    def _serve_gzipped(self, file_path, content_type):
        """Serve a file with gzip compression and aggressive caching."""
        mtime = file_path.stat().st_mtime
        cache_key = str(file_path)

        # Check gzip cache (avoids re-compressing on every request)
        cached = self._gzip_cache.get(cache_key)
        if cached and cached[0] == mtime:
            compressed = cached[1]
        else:
            with open(str(file_path), "rb") as f:
                raw = f.read()
            compressed = gzip.compress(raw, compresslevel=6)
            self._gzip_cache[cache_key] = (mtime, compressed)

        self.send_response(200)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Encoding", "gzip")
        self.send_header("Content-Length", str(len(compressed)))
        # Cache static assets aggressively (reload via Ctrl+Shift+R if needed)
        self.send_header("Cache-Control", "public, max-age=3600, stale-while-revalidate=86400")
        self.end_headers()
        self.wfile.write(compressed)

    def do_GET(self):
        """Serve graph_data.json with gzip; other files with caching."""
        clean_path = self.path.split("?")[0]

        if clean_path in ("/graph_data.json", "/graph_clusters.json"):
            candidates = [
                VIS_DIR / "graph_clusters.json",
                VIS_DIR / "graph_data.json",
                VIS_DIR_V9 / "graph_data.json",
                VIS_DIR_V5 / "graph_data.json",
            ]
            gj = next((c for c in candidates if c.exists()), None)
            if not gj:
                self.send_error(404, "graph data not found in V15/V9/V5")
                return
            # Gzip compress: 4.3 MB → ~400 KB
            accept = self.headers.get("Accept-Encoding", "")
            if "gzip" in accept:
                self._serve_gzipped(gj, "application/json")
            else:
                # Fallback: uncompressed chunked
                self.send_response(200)
                self.send_header("Content-Type", "application/json")
                sz = gj.stat().st_size
                self.send_header("Content-Length", str(sz))
                self.send_header("Cache-Control", "public, max-age=3600")
                self.end_headers()
                with open(str(gj), "rb") as f:
                    while True:
                        chunk = f.read(65536)
                        if not chunk:
                            break
                        self.wfile.write(chunk)
            return

        # For HTML/JS/CSS: add caching headers
        ext = os.path.splitext(clean_path)[1].lower()
        if ext in (".html", ".js", ".css") and "gzip" in self.headers.get("Accept-Encoding", ""):
            fpath = VIS_DIR / clean_path.lstrip("/")
            if fpath.exists():
                ct = self.guess_type(clean_path)
                self._serve_gzipped(fpath, ct)
                return

        super().do_GET()


class MBPServer(threading.Thread):
    """HTTP server thread for static assets."""

    def __init__(self, port=0):
        super().__init__(daemon=True, name="mbp-http")
        self.port = port or _find_free_port()
        self.server = http.server.HTTPServer(("127.0.0.1", self.port), _QuietHandler)

    def run(self):
        self.server.serve_forever()

    def shutdown(self):
        self.server.shutdown()


# ---------------------------------------------------------------------------
# Banner
# ---------------------------------------------------------------------------
def _print_banner(port, stats):
    sep = _sep_days()
    nodes = stats.get("node_count", "?")
    edges = stats.get("edge_count", "?")
    chains = stats.get("chain_count", "?")
    gaps = stats.get("gap_open", "?")
    ver = stats.get("brain_version", "?")
    banner = (
        "\n"
        "+==============================================================+\n"
        f"|       THEMANBEARPIG v{VERSION} -- SINGULARITY CONVERGENCE      |\n"
        "+==============================================================+\n"
        f"|  HTTP  : http://127.0.0.1:{port:<5}                             |\n"
        f"|  Brain : {str(BRAIN_DB):<50} |\n"
        f"|  Lit DB: {str(LIT_DB):<50} |\n"
        f"|  Version: {ver:<4}  Nodes: {nodes:<8}  Edges: {edges:<8}         |\n"
        f"|  Chains: {chains:<6}  Open Gaps: {gaps:<6}                        |\n"
        f"|  SEPARATION: {sep} DAYS since July 29, 2025              |\n"
        "+==============================================================+\n"
    )
    print(banner)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(
        description=f"THEMANBEARPIG v{VERSION} -- SINGULARITY CONVERGENCE"
    )
    parser.add_argument("--debug", action="store_true", help="Enable developer tools")
    parser.add_argument("--port", type=int, default=0, help="HTTP server port (0=auto)")
    parser.add_argument("--hunt", action="store_true", help="Launch KRAKEN hunting on startup")
    parser.add_argument("--export", action="store_true", help="Re-export graph data before launch")
    parser.add_argument("--rounds", type=int, default=3, help="KRAKEN hunting rounds (with --hunt)")
    parser.add_argument("--count", type=int, default=10, help="Files per KRAKEN round (with --hunt)")
    parser.add_argument("--focus", default="all", help="KRAKEN focus mode (with --hunt)")
    parser.add_argument("--v15", action="store_true", help="Use V15 visualization (legacy)")
    parser.add_argument("--v7", action="store_true", default=True, help="Use V7 SELFEVOLVE visualization (default)")
    args = parser.parse_args()

    try:
        import webview
    except ImportError:
        print("ERROR: pywebview not installed. Run: pip install pywebview")
        sys.exit(1)

    if not BRAIN_DB.exists() and not LIT_DB.exists():
        print("WARNING: Neither mbp_brain.db nor litigation_context.db found.")
        print("The application will run in limited mode.")

    # V7 SELFEVOLVE CONVERGENCE — select visualization
    if args.v15:
        active_vis = VIS_DIR  # V15 legacy
        vis_label = "V15 (legacy)"
    elif VIS_DIR_V7.exists() and (VIS_DIR_V7 / "index.html").exists():
        active_vis = VIS_DIR_V7  # V7 SELFEVOLVE (default)
        vis_label = "V7 SELFEVOLVE"
    elif VIS_DIR.exists() and (VIS_DIR / "index.html").exists():
        active_vis = VIS_DIR  # Fallback to V15
        vis_label = "V15 (fallback)"
    else:
        active_vis = VIS_DIR
        vis_label = "V15"

    if not active_vis.exists():
        active_vis.mkdir(parents=True, exist_ok=True)

    index_html = active_vis / "index.html"
    if not index_html.exists():
        print(f"ERROR: {index_html} not found.")
        print(f"Active visualization: {vis_label} at {active_vis}")
        sys.exit(1)

    print(f"Visualization: {vis_label}")

    if args.export and EXPORT_SCRIPT.exists():
        print("Re-exporting graph data...")
        subprocess.run(
            [_PYTHON, "-I", str(EXPORT_SCRIPT)],
            cwd=str(REPO_ROOT), timeout=300,
        )

    api = UnifiedAPI()
    stats = api.get_stats()

    # Override VIS_DIR for HTTP server
    VIS_DIR = active_vis
    server = MBPServer(port=args.port)
    server.start()

    _print_banner(server.port, stats)

    if args.hunt:
        print(f"Starting KRAKEN hunter: {args.rounds} rounds, {args.count} files, focus={args.focus}")
        api.start_kraken(rounds=args.rounds, count=args.count, focus=args.focus)

    sep = _sep_days()
    title = f"THEMANBEARPIG v{VERSION} [{vis_label}] -- {sep} Days Separated"

    window = webview.create_window(
        title,
        url=f"http://127.0.0.1:{server.port}/index.html",
        js_api=api,
        width=APP_WIDTH,
        height=APP_HEIGHT,
        min_size=(APP_MIN_W, APP_MIN_H),
        background_color=APP_BG,
        text_select=True,
    )

    webview.start(debug=args.debug)
    server.shutdown()
    print("THEMANBEARPIG exited.")


if __name__ == "__main__":
    import io as _io
    try:
        sys.stdout = _io.TextIOWrapper(
            sys.stdout.buffer, encoding="utf-8", errors="replace", line_buffering=True
        )
    except Exception:
        pass
    main()
