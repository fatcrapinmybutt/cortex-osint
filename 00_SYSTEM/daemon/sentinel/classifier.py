"""SENTINEL File Classifier — MEEK lane routing + category detection.

Extracts content from PDF/DOCX/TXT/CSV/JSON/MD/HTML files, classifies
into litigation lanes (A-F) and categories (20 types), and returns
a structured classification result.
"""
from __future__ import annotations

import os
import re
import hashlib
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

logger = logging.getLogger("sentinel.classifier")

# ---------------------------------------------------------------------------
# MEEK Lane Patterns — compiled regex, priority order E→D→F→A→B, default C
# Mirrors 00_SYSTEM/engines/bridge/meek.py
# ---------------------------------------------------------------------------
LANE_PRIORITY: list[str] = ["E", "D", "F", "A", "B"]

MEEK_PATTERNS: dict[str, re.Pattern[str]] = {
    "E": re.compile(
        r"McNeill|judicial|bias|JTC|canon|misconduct|benchbook|ex\s*parte"
        r"|disqualif|recus|tenure|violation",
        re.IGNORECASE,
    ),
    "D": re.compile(
        r"PPO|protection\s*order|5907|stalking|harassment"
        r"|personal\s*protection|domestic\s*violence",
        re.IGNORECASE,
    ),
    "F": re.compile(
        r"COA|366810|appeal|appellant|appellee|appendix"
        r"|court\s*of\s*appeals|MSC|superintending",
        re.IGNORECASE,
    ),
    "A": re.compile(
        r"custody|parenting|001507|Watson|child|visitation|FOC"
        r"|best\s*interest|MCL\s*722|guardianship",
        re.IGNORECASE,
    ),
    "B": re.compile(
        r"Shady\s*Oaks|eviction|housing|trailer|002760|habitability"
        r"|landlord|tenant|lease|rent",
        re.IGNORECASE,
    ),
}

# ---------------------------------------------------------------------------
# Category Patterns — keyword regex for 20 evidence categories
# ---------------------------------------------------------------------------
CATEGORY_PATTERNS: dict[str, re.Pattern[str]] = {
    "court_filing": re.compile(
        r"\b(?:motion|order|judgment|complaint|brief|petition|stipulation"
        r"|plea|arraignment|sentence|verdict|ruling|summons)\b",
        re.IGNORECASE,
    ),
    "evidence": re.compile(
        r"\b(?:exhibit|affidavit|declaration|testimony|deposition"
        r"|sworn\s*statement)\b",
        re.IGNORECASE,
    ),
    "authority": re.compile(
        r"\b(?:MCR|MCL|MRE|statute|case\s*law|court\s*rule"
        r"|USC|FRCP|precedent)\b",
        re.IGNORECASE,
    ),
    "analysis": re.compile(
        r"\b(?:analysis|report|summary|assessment|evaluation"
        r"|findings|memorandum|memo)\b",
        re.IGNORECASE,
    ),
    "police": re.compile(
        r"\b(?:police|incident|officer|NSPD|report\s*number"
        r"|patrol|arrest|dispatch|badge)\b",
        re.IGNORECASE,
    ),
    "financial": re.compile(
        r"\b(?:bank|statement|income|tax|W-2|1099|pay\s*stub"
        r"|ledger|invoice|receipt|expense)\b",
        re.IGNORECASE,
    ),
    "medical": re.compile(
        r"\b(?:medical|doctor|hospital|prescription|diagnosis"
        r"|HealthWest|therapy|counseling|mental\s*health|LOCUS)\b",
        re.IGNORECASE,
    ),
    "correspondence": re.compile(
        r"\b(?:letter|email|notice|correspondence|message"
        r"|communication|fax|memo\s*to)\b",
        re.IGNORECASE,
    ),
    "media": re.compile(
        r"\b(?:photo|video|audio|recording|screenshot|image"
        r"|camera|surveillance)\b",
        re.IGNORECASE,
    ),
    "government": re.compile(
        r"\b(?:FOIA|government|agency|DHHS|CPS|state\s*of"
        r"|federal|department)\b",
        re.IGNORECASE,
    ),
    "housing": re.compile(
        r"\b(?:lease|rental|property|deed|title|mortgage"
        r"|landlord|tenant|mobile\s*home)\b",
        re.IGNORECASE,
    ),
    "employment": re.compile(
        r"\b(?:employment|employer|fired|terminated|wages"
        r"|job|work|salary|hire|resign)\b",
        re.IGNORECASE,
    ),
    "education": re.compile(
        r"\b(?:school|education|student|teacher|IEP"
        r"|enrollment|transcript|diploma)\b",
        re.IGNORECASE,
    ),
    "app_export": re.compile(
        r"\b(?:AppClose|OurFamilyWizard|TalkingParents|app\s*export"
        r"|message\s*log|chat\s*export)\b",
        re.IGNORECASE,
    ),
    "transcript": re.compile(
        r"\b(?:transcript|hearing\s*transcript|court\s*reporter"
        r"|verbatim|proceedings)\b",
        re.IGNORECASE,
    ),
    "forensic": re.compile(
        r"\b(?:forensic|chain\s*of\s*custody|hash|metadata"
        r"|digital\s*evidence|extraction)\b",
        re.IGNORECASE,
    ),
}

# Map category → canonical subdirectory under LitigationOS root
CATEGORY_DIRS: dict[str, str] = {
    "court_filing": "03_COURT",
    "evidence": "01_EVIDENCE",
    "authority": "02_AUTHORITY",
    "analysis": "04_ANALYSIS",
    "police": "01_EVIDENCE/police",
    "financial": "01_EVIDENCE/financial",
    "medical": "01_EVIDENCE/medical",
    "correspondence": "01_EVIDENCE/correspondence",
    "media": "08_MEDIA",
    "government": "01_EVIDENCE/government",
    "housing": "01_EVIDENCE/housing",
    "employment": "01_EVIDENCE/employment",
    "education": "01_EVIDENCE/education",
    "app_export": "01_EVIDENCE/app_export",
    "transcript": "01_EVIDENCE/transcript",
    "forensic": "01_EVIDENCE/forensic",
    "unknown": "12_WORKSPACE/unsorted",
}

# Supported file extensions for processing
SUPPORTED_EXTENSIONS: frozenset[str] = frozenset({
    ".pdf", ".docx", ".txt", ".csv", ".json", ".md", ".html", ".htm",
    ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff",
    ".mp3", ".wav", ".mp4", ".avi", ".mov",
})

# Ignore patterns (glob-style, checked by name)
IGNORE_PATTERNS: tuple[re.Pattern[str], ...] = (
    re.compile(r".*\.tmp$", re.IGNORECASE),
    re.compile(r".*\.lock$", re.IGNORECASE),
    re.compile(r".*\.partial$", re.IGNORECASE),
    re.compile(r"^~\$.*"),
    re.compile(r"__pycache__"),
    re.compile(r".*\.pyc$", re.IGNORECASE),
    re.compile(r".*\.db-shm$", re.IGNORECASE),
    re.compile(r".*\.db-wal$", re.IGNORECASE),
)

# Max bytes to read for classification (first 2000 chars)
MAX_EXTRACT_CHARS: int = 2000


@dataclass
class ClassificationResult:
    """Result of classifying a single file."""
    file_path: str
    lane: str
    category: str
    confidence: float
    summary: str
    sha256: str = ""
    file_size: int = 0
    all_lanes: list[str] = field(default_factory=list)
    all_categories: list[str] = field(default_factory=list)

    @property
    def canonical_dir(self) -> str:
        """Return the canonical subdirectory for this classification."""
        base = CATEGORY_DIRS.get(self.category, CATEGORY_DIRS["unknown"])
        if self.category not in ("authority", "analysis", "media", "unknown"):
            return f"{base}/{self.lane}"
        return base


class FileClassifier:
    """Classifies files by litigation lane and evidence category.

    Uses MEEK compiled regex for lane detection and keyword patterns
    for category classification. Extracts content from PDF, DOCX, and
    text-based formats.
    """

    def __init__(self) -> None:
        self._pdf_available = self._check_pypdfium2()
        self._docx_available = self._check_python_docx()

    @staticmethod
    def _check_pypdfium2() -> bool:
        try:
            import pypdfium2  # noqa: F401
            return True
        except ImportError:
            logger.info("pypdfium2 not available; PDF classification will use raw text fallback")
            return False

    @staticmethod
    def _check_python_docx() -> bool:
        try:
            import docx  # noqa: F401
            return True
        except ImportError:
            logger.info("python-docx not available; DOCX classification disabled")
            return False

    def should_process(self, file_path: str) -> bool:
        """Check if a file should be processed based on extension and ignore rules."""
        name = os.path.basename(file_path)
        for pat in IGNORE_PATTERNS:
            if pat.search(name):
                return False
        ext = os.path.splitext(name)[1].lower()
        return ext in SUPPORTED_EXTENSIONS

    def classify(self, file_path: str) -> ClassificationResult:
        """Classify a file into a litigation lane and category.

        Args:
            file_path: Absolute path to the file.

        Returns:
            ClassificationResult with lane, category, confidence, and summary.
        """
        path = Path(file_path)
        if not path.exists():
            return ClassificationResult(
                file_path=file_path, lane="C", category="unknown",
                confidence=0.0, summary="File not found",
            )

        file_size = path.stat().st_size
        sha256 = self._compute_sha256(file_path)
        text = self._extract_text(file_path)

        # Classify lane
        lane = self._classify_lane(text, path.name)
        all_lanes = self._classify_lanes_multi(text, path.name)

        # Classify category
        category, cat_confidence, all_cats = self._classify_category(text, path.name)

        # Build summary (first 120 chars of extracted text, cleaned)
        summary = self._build_summary(text, path.name)

        return ClassificationResult(
            file_path=file_path,
            lane=lane,
            category=category,
            confidence=cat_confidence,
            summary=summary,
            sha256=sha256,
            file_size=file_size,
            all_lanes=all_lanes,
            all_categories=all_cats,
        )

    def _classify_lane(self, text: str, filename: str) -> str:
        """Assign primary lane using MEEK priority order."""
        combined = f"{filename} {text}"
        for lane_id in LANE_PRIORITY:
            if MEEK_PATTERNS[lane_id].search(combined):
                return lane_id
        return "C"

    def _classify_lanes_multi(self, text: str, filename: str) -> list[str]:
        """Return all matching lanes in priority order."""
        combined = f"{filename} {text}"
        matches = [lid for lid in LANE_PRIORITY if MEEK_PATTERNS[lid].search(combined)]
        return matches if matches else ["C"]

    def _classify_category(
        self, text: str, filename: str
    ) -> tuple[str, float, list[str]]:
        """Assign category and confidence based on keyword density.

        Returns:
            Tuple of (best_category, confidence, all_matching_categories).
        """
        combined = f"{filename} {text}"
        scores: dict[str, int] = {}
        for cat_name, pattern in CATEGORY_PATTERNS.items():
            hits = pattern.findall(combined)
            if hits:
                scores[cat_name] = len(hits)

        if not scores:
            return "unknown", 0.1, ["unknown"]

        total_hits = sum(scores.values())
        best_cat = max(scores, key=scores.get)  # type: ignore[arg-type]
        best_hits = scores[best_cat]

        # Confidence: ratio of best category hits to total, scaled
        raw_conf = best_hits / max(total_hits, 1)
        # Boost confidence if there are many hits overall
        density_boost = min(total_hits / 20.0, 0.3)
        confidence = min(raw_conf + density_boost, 1.0)

        all_cats = sorted(scores, key=scores.get, reverse=True)  # type: ignore[arg-type]
        return best_cat, round(confidence, 3), all_cats

    def _extract_text(self, file_path: str) -> str:
        """Extract text content from a file (first MAX_EXTRACT_CHARS chars)."""
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            return self._extract_pdf(file_path)
        if ext == ".docx":
            return self._extract_docx(file_path)
        if ext in (".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff",
                    ".mp3", ".wav", ".mp4", ".avi", ".mov"):
            # Media files: classify by filename only
            return ""
        # Text-based formats
        return self._extract_text_file(file_path)

    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF using pypdfium2, fallback to raw bytes."""
        if self._pdf_available:
            try:
                import pypdfium2 as pdfium
                pdf = pdfium.PdfDocument(file_path)
                parts: list[str] = []
                chars_collected = 0
                for i in range(len(pdf)):
                    if chars_collected >= MAX_EXTRACT_CHARS:
                        break
                    page = pdf[i]
                    text = page.get_textpage().get_text_range()
                    parts.append(text)
                    chars_collected += len(text)
                    page.close()
                pdf.close()
                return " ".join(parts)[:MAX_EXTRACT_CHARS]
            except Exception as exc:
                logger.warning("pypdfium2 failed on %s: %s — using raw fallback", file_path, exc)

        # Fallback: read raw bytes and extract printable ASCII
        try:
            with open(file_path, "rb") as f:
                raw = f.read(8192)
            # Extract printable ASCII sequences of 4+ chars
            printable = re.findall(rb"[\x20-\x7e]{4,}", raw)
            return " ".join(p.decode("ascii", errors="ignore") for p in printable)[:MAX_EXTRACT_CHARS]
        except Exception as exc:
            logger.warning("Raw PDF fallback failed on %s: %s", file_path, exc)
            return ""

    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX using python-docx."""
        if not self._docx_available:
            return ""
        try:
            import docx
            doc = docx.Document(file_path)
            parts: list[str] = []
            chars_collected = 0
            for para in doc.paragraphs:
                if chars_collected >= MAX_EXTRACT_CHARS:
                    break
                text = para.text.strip()
                if text:
                    parts.append(text)
                    chars_collected += len(text)
            return " ".join(parts)[:MAX_EXTRACT_CHARS]
        except Exception as exc:
            logger.warning("DOCX extraction failed on %s: %s", file_path, exc)
            return ""

    def _extract_text_file(self, file_path: str) -> str:
        """Read text from TXT/CSV/JSON/MD/HTML with encoding detection."""
        encodings = ("utf-8", "cp1252", "latin-1")
        for enc in encodings:
            try:
                with open(file_path, "r", encoding=enc, errors="strict") as f:
                    return f.read(MAX_EXTRACT_CHARS)
            except (UnicodeDecodeError, UnicodeError):
                continue
            except Exception as exc:
                logger.warning("Text read failed on %s with %s: %s", file_path, enc, exc)
                return ""
        # Last resort: lossy read
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return f.read(MAX_EXTRACT_CHARS)
        except Exception:
            return ""

    @staticmethod
    def _compute_sha256(file_path: str) -> str:
        """Compute SHA-256 hash of a file (streaming, 1MB chunks)."""
        sha = hashlib.sha256()
        try:
            with open(file_path, "rb") as f:
                while True:
                    chunk = f.read(1024 * 1024)
                    if not chunk:
                        break
                    sha.update(chunk)
            return sha.hexdigest()
        except Exception as exc:
            logger.warning("SHA-256 failed on %s: %s", file_path, exc)
            return ""

    @staticmethod
    def _build_summary(text: str, filename: str) -> str:
        """Build a brief human-readable summary for logging."""
        if not text:
            return f"[{filename}] No text extracted"
        clean = re.sub(r"\s+", " ", text).strip()
        snippet = clean[:120]
        if len(clean) > 120:
            snippet += "..."
        return snippet


if __name__ == "__main__":
    import sys
    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

    if len(sys.argv) < 2:
        print("Usage: python -I classifier.py <file_path> [file_path ...]")
        sys.exit(1)

    clf = FileClassifier()
    for fpath in sys.argv[1:]:
        if not os.path.isfile(fpath):
            print(f"  SKIP (not a file): {fpath}")
            continue
        result = clf.classify(fpath)
        print(f"  File: {result.file_path}")
        print(f"  Lane: {result.lane} (all: {result.all_lanes})")
        print(f"  Category: {result.category} (confidence: {result.confidence})")
        print(f"  SHA-256: {result.sha256[:16]}...")
        print(f"  Summary: {result.summary}")
        print()
